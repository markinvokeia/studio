#!/usr/bin/env python3
"""
Invoke IA Documentation DOCX Styler
Generates a modern, styled DOCX from the markdown guide.

Usage:
    python3 docs/style_docx.py
"""
import os
import subprocess
import urllib.request
from copy import deepcopy

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ──────────────────────────────────────────────────────────────────
DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
MD_FILE   = os.path.join(DOCS_DIR, 'guia-uso-invoke-ia.md')
DOCX_BASE = os.path.join(DOCS_DIR, 'guia-uso-invoke-ia-base.docx')
DOCX_OUT  = os.path.join(DOCS_DIR, 'guia-uso-invoke-ia.docx')
LOGO_WEBP = '/tmp/invokeia_logo.webp'
LOGO_PNG  = '/tmp/invokeia_logo.png'
LOGO_URL  = 'https://www.invokeia.com/assets/InvokeIA_C@4x-4T0dztu0.webp'

# ── Brand Colors ───────────────────────────────────────────────────────────
# Primary: hsl(263, 67%, 35%) ≈ #4B1D97 (deep violet-purple)
PRIMARY       = RGBColor(0x4B, 0x1D, 0x97)
# Accent: lighter purple
ACCENT        = RGBColor(0x7C, 0x3A, 0xED)
# Light tint for callout backgrounds
CALLOUT_BG    = RGBColor(0xEE, 0xE6, 0xFB)
# Dark body text
DARK          = RGBColor(0x1F, 0x29, 0x37)
# Medium gray for secondary text
GRAY          = RGBColor(0x6B, 0x72, 0x80)
# Table header background
TABLE_HEADER  = RGBColor(0x4B, 0x1D, 0x97)
# Alternating row background
TABLE_ALT     = RGBColor(0xF5, 0xF3, 0xFF)
# Table border color
BORDER        = RGBColor(0xD1, 0xD5, 0xDB)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)

# ── Fonts ──────────────────────────────────────────────────────────────────
BODY_FONT    = 'Calibri'
HEADING_FONT = 'Calibri Light'


# ── Helpers ────────────────────────────────────────────────────────────────

def rgb_hex(rgb: RGBColor) -> str:
    """Return RRGGBB hex string from an RGBColor (tuple subclass; str() returns 6-char hex)."""
    return str(rgb)  # e.g. '4B1D97'


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Apply borders to a table cell. Colors are RGBColor instances."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if color is not None:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '6')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), rgb_hex(color))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_background(cell, rgb: RGBColor):
    """Fill cell with a background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rgb_hex(rgb))
    tcPr.append(shd)


def add_horizontal_rule(doc):
    """Insert a colored horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), rgb_hex(PRIMARY))
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_page_number_field(paragraph):
    """Insert PAGE / NUMPAGES field into a paragraph."""
    run = paragraph.add_run()
    fldBegin = OxmlElement('w:fldChar')
    fldBegin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldBegin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run._r.append(instrText)

    fldSep = OxmlElement('w:fldChar')
    fldSep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldSep)

    fldEnd = OxmlElement('w:fldChar')
    fldEnd.set(qn('w:fldCharType'), 'end')
    run._r.append(fldEnd)

    paragraph.add_run(' / ')

    run2 = paragraph.add_run()
    fldBegin2 = OxmlElement('w:fldChar')
    fldBegin2.set(qn('w:fldCharType'), 'begin')
    run2._r.append(fldBegin2)

    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = ' NUMPAGES '
    run2._r.append(instrText2)

    fldSep2 = OxmlElement('w:fldChar')
    fldSep2.set(qn('w:fldCharType'), 'separate')
    run2._r.append(fldSep2)

    fldEnd2 = OxmlElement('w:fldChar')
    fldEnd2.set(qn('w:fldCharType'), 'end')
    run2._r.append(fldEnd2)


def download_logo():
    """Download logo webp and convert to PNG with sips."""
    print('Downloading logo...')
    try:
        urllib.request.urlretrieve(LOGO_URL, LOGO_WEBP)
        print(f'  → saved to {LOGO_WEBP}')
    except Exception as e:
        print(f'  ✗ Download failed: {e}')
        return False

    print('Converting webp → PNG...')
    result = subprocess.run(
        ['sips', '-s', 'format', 'png', LOGO_WEBP, '--out', LOGO_PNG],
        capture_output=True, text=True
    )
    if result.returncode != 0:
        print(f'  ✗ Conversion failed: {result.stderr}')
        return False
    print(f'  → saved to {LOGO_PNG}')
    return True


def generate_base_docx():
    """Run pandoc to generate a clean base DOCX from the markdown."""
    print('Generating base DOCX with pandoc...')
    result = subprocess.run(
        ['/opt/homebrew/bin/pandoc',
         MD_FILE,
         '-o', DOCX_BASE,
         '--wrap=none',
         '-f', 'markdown',
         '-t', 'docx'],
        capture_output=True, text=True
    )
    if result.returncode != 0:
        print(f'  ✗ pandoc failed: {result.stderr}')
        return False
    print(f'  → {DOCX_BASE}')
    return True


# ── Style application ──────────────────────────────────────────────────────

def apply_document_defaults(doc):
    """Set default font and size for the whole document."""
    styles = doc.styles

    # Normal / body
    normal = styles['Normal']
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = Pt(14)

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = HEADING_FONT
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = PRIMARY
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = HEADING_FONT
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = PRIMARY
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = HEADING_FONT
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = ACCENT
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    # Heading 4
    try:
        h4 = styles['Heading 4']
        h4.font.name = BODY_FONT
        h4.font.size = Pt(11)
        h4.font.bold = True
        h4.font.color.rgb = DARK
        h4.paragraph_format.space_before = Pt(8)
        h4.paragraph_format.space_after = Pt(3)
    except KeyError:
        pass

    # Block Quote (used for metadata callouts)
    try:
        bq = styles['Block Text']
        bq.font.name = BODY_FONT
        bq.font.size = Pt(9.5)
        bq.font.color.rgb = GRAY
    except KeyError:
        pass

    # List Bullet
    try:
        lb = styles['List Bullet']
        lb.font.name = BODY_FONT
        lb.font.size = Pt(10.5)
    except KeyError:
        pass

    # List Number
    try:
        ln = styles['List Number']
        ln.font.name = BODY_FONT
        ln.font.size = Pt(10.5)
    except KeyError:
        pass

    # Code / Verbatim
    for style_name in ('Verbatim Char', 'Code', 'Verbatim'):
        try:
            code_s = styles[style_name]
            code_s.font.name = 'Courier New'
            code_s.font.size = Pt(9)
            code_s.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        except KeyError:
            pass


def style_all_tables(doc):
    """Apply styled borders and header formatting to every table."""
    header_bg = TABLE_HEADER
    alt_bg    = TABLE_ALT
    for table in doc.tables:
        # Use available table style (pandoc generates 'Table')
        try:
            table.style = 'Table Grid'
        except KeyError:
            pass  # style not available; borders applied manually below
        for row_idx, row in enumerate(table.rows):
            is_header = (row_idx == 0)
            for cell in row.cells:
                # Borders on all 4 sides
                set_cell_border(cell,
                    top=BORDER, bottom=BORDER,
                    left=BORDER, right=BORDER)
                # Header row: purple background, white bold text
                if is_header:
                    set_cell_background(cell, header_bg)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.color.rgb = WHITE
                            run.font.bold = True
                            run.font.name = BODY_FONT
                            run.font.size = Pt(10)
                        para.paragraph_format.space_before = Pt(3)
                        para.paragraph_format.space_after = Pt(3)
                # Alternating rows: very light tint
                elif row_idx % 2 == 0:
                    set_cell_background(cell, alt_bg)

                # Ensure correct font in all data cells
                if not is_header:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if not run.font.name:
                                run.font.name = BODY_FONT
                            run.font.size = Pt(10)


def add_header_with_logo(doc, logo_path):
    """Add a header with the logo on the left and document title on the right."""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        # Clear default content
        for p in header.paragraphs:
            p.clear()

        # Use the first paragraph
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.clear()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # Logo run on the left
        if os.path.exists(logo_path):
            run_logo = p.add_run()
            run_logo.add_picture(logo_path, height=Inches(0.35))

        # Spacer
        p.add_run('\t')

        # Title run on the right
        run_title = p.add_run('Guía de Uso — Invoke IA')
        run_title.font.name = HEADING_FONT
        run_title.font.size = Pt(9)
        run_title.font.color.rgb = GRAY

        # Tab stop for right-align
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '9072')   # ~6.3 inches (A4 width - margins)
        tabs.append(tab)
        pPr.append(tabs)

        # Separator line below header
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), rgb_hex(ACCENT))
        pBdr.append(bottom)
        pPr.append(pBdr)


def add_footer_with_logo(doc, logo_path):
    """Add a footer with logo + name on the left and page number on the right."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False

        for p in footer.paragraphs:
            p.clear()

        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # Logo tiny
        if os.path.exists(logo_path):
            run_logo = p.add_run()
            run_logo.add_picture(logo_path, height=Inches(0.22))

        # Name
        run_name = p.add_run('  Invoke IA')
        run_name.font.name = BODY_FONT
        run_name.font.size = Pt(8)
        run_name.font.color.rgb = GRAY

        # Tab to right
        p.add_run('\t')

        # Page number
        run_pg = p.add_run('Página ')
        run_pg.font.name = BODY_FONT
        run_pg.font.size = Pt(8)
        run_pg.font.color.rgb = GRAY

        add_page_number_field(p)

        # Tab stop
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '9072')
        tabs.append(tab)
        pPr.append(tabs)

        # Top separator line
        pBdr = OxmlElement('w:pBdr')
        top_el = OxmlElement('w:top')
        top_el.set(qn('w:val'), 'single')
        top_el.set(qn('w:sz'), '4')
        top_el.set(qn('w:space'), '1')
        top_el.set(qn('w:color'), rgb_hex(ACCENT))
        pBdr.append(top_el)
        pPr.append(pBdr)


def add_title_page(doc, logo_path):
    """Prepend a branded title page before the document content."""
    # We'll insert paragraphs at position 0 using XML manipulation
    body = doc.element.body

    def make_para(text, font_name, font_size, bold=False, color=None, align='center', space_before=0, space_after=8):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), align)
        pPr.append(jc)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), str(int(space_before * 20)))
        spacing.set(qn('w:after'), str(int(space_after * 20)))
        pPr.append(spacing)
        p.append(pPr)

        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(font_size * 2)))
        rPr.append(szCs)
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        if color:
            clr = OxmlElement('w:color')
            clr.set(qn('w:val'), rgb_hex(color))
            rPr.append(clr)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p.append(r)
        return p

    def make_page_break():
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        r.append(br)
        p.append(r)
        return p

    def make_logo_para(logo_path):
        """Create a paragraph with a centered logo image."""
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), str(int(180 * 20)))
        spacing.set(qn('w:after'), str(int(20 * 20)))
        pPr.append(spacing)
        p.append(pPr)
        return p  # logo will be added via Document paragraph

    # Build title page elements (insert in reverse since we prepend)
    title_elements = []

    # Spacer at top (simulate vertical centering with spacing)
    title_elements.append(make_para('', HEADING_FONT, 12, space_before=20, space_after=0))

    # Version and date line (light, below divider)
    title_elements.append(make_para('versión 3.2 · mayo 2026', BODY_FONT, 11,
                                    color=GRAY, align='center', space_before=0, space_after=6))

    # Subtitle
    title_elements.append(make_para(
        'Sistema de Gestión para Clínicas Odontológicas y de Salud',
        HEADING_FONT, 14, color=ACCENT, align='center', space_before=4, space_after=8
    ))

    # Main title
    title_elements.append(make_para('Guía de Uso', HEADING_FONT, 36,
                                    bold=True, color=PRIMARY, align='center',
                                    space_before=16, space_after=4))

    # Short divider line (empty para with bottom border)
    div = OxmlElement('w:p')
    divPPr = OxmlElement('w:pPr')
    divJc = OxmlElement('w:jc')
    divJc.set(qn('w:val'), 'center')
    divPPr.append(divJc)
    divSpacing = OxmlElement('w:spacing')
    divSpacing.set(qn('w:before'), '120')
    divSpacing.set(qn('w:after'), '120')
    divPPr.append(divSpacing)
    divBdr = OxmlElement('w:pBdr')
    divBottom = OxmlElement('w:bottom')
    divBottom.set(qn('w:val'), 'single')
    divBottom.set(qn('w:sz'), '12')
    divBottom.set(qn('w:space'), '1')
    divBottom.set(qn('w:color'), rgb_hex(PRIMARY))
    divBdr.append(divBottom)
    divPPr.append(divBdr)
    div.append(divPPr)
    title_elements.append(div)

    # Page break at end of title page
    title_elements.append(make_page_break())

    # Insert title elements at beginning of body (before first existing element)
    first_child = body[0]
    for el in title_elements:
        body.insert(0, el)

    # Add logo image using docx Document API (trickier to do via raw XML)
    # We'll add it properly by finding the spacer paragraph at position 0 and replacing it
    if os.path.exists(logo_path):
        # Find the first paragraph we inserted (currently index 0 since we prepended in order)
        # and add the logo to it
        for i, para in enumerate(doc.paragraphs[:10]):
            if para.text == '' and i < 5:
                # Add logo to this empty paragraph
                run = para.add_run()
                try:
                    run.add_picture(logo_path, width=Inches(3.0))
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f'  ⚠ Could not add logo to title page: {e}')
                break


def style_blockquotes(doc):
    """Style blockquote paragraphs (metadata callouts) with left border and light background."""
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        # Pandoc uses 'Block Text' for blockquotes
        if 'Block' in style_name or 'Quote' in style_name:
            pPr = para._p.get_or_add_pPr()

            # Left bar border (like a callout)
            pBdr = OxmlElement('w:pBdr')
            left_el = OxmlElement('w:left')
            left_el.set(qn('w:val'), 'single')
            left_el.set(qn('w:sz'), '12')
            left_el.set(qn('w:space'), '4')
            left_el.set(qn('w:color'), rgb_hex(ACCENT))
            pBdr.append(left_el)
            pPr.append(pBdr)

            # Light background shading
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), rgb_hex(CALLOUT_BG))
            pPr.append(shd)

            # Indent and spacing
            ind = OxmlElement('w:ind')
            ind.set(qn('w:left'), '360')
            ind.set(qn('w:right'), '360')
            pPr.append(ind)
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after = Pt(4)

            # Style the text
            for run in para.runs:
                run.font.name = BODY_FONT
                run.font.size = Pt(9.5)
                if not run.font.color.type:
                    run.font.color.rgb = GRAY


def fix_code_spans(doc):
    """Style inline code spans with a monospace font and distinct color."""
    for para in doc.paragraphs:
        for run in para.runs:
            if run.font.name in ('Courier New', 'Courier', 'Lucida Console', 'Consolas'):
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xB0, 0x3A, 0x2E)   # dark red


def set_page_margins(doc):
    """Set comfortable A4 margins."""
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.8)
        section.bottom_margin = Cm(2.5)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.2)


# ── Main ───────────────────────────────────────────────────────────────────

def main():
    print('=== Invoke IA DOCX Styler ===\n')

    # 1. Download & convert logo
    logo_ok = download_logo()
    logo_path = LOGO_PNG if logo_ok and os.path.exists(LOGO_PNG) else None

    # 2. Generate base DOCX with pandoc
    if not generate_base_docx():
        return

    # 3. Open and style
    print('Styling document...')
    doc = Document(DOCX_BASE)

    # Set page size & margins
    set_page_margins(doc)
    print('  ✓ Page layout')

    # Apply default styles
    apply_document_defaults(doc)
    print('  ✓ Typography')

    # Style all tables
    style_all_tables(doc)
    print('  ✓ Tables')

    # Style blockquote callouts
    style_blockquotes(doc)
    print('  ✓ Callout blocks')

    # Style inline code
    fix_code_spans(doc)
    print('  ✓ Code spans')

    # Add header with logo
    add_header_with_logo(doc, logo_path or '')
    print('  ✓ Header')

    # Add footer with page numbers
    add_footer_with_logo(doc, logo_path or '')
    print('  ✓ Footer')

    # Add branded title page
    add_title_page(doc, logo_path or '')
    print('  ✓ Title page')

    # 4. Save
    doc.save(DOCX_OUT)
    print(f'\n✓ Saved: {DOCX_OUT}')
    size_kb = os.path.getsize(DOCX_OUT) // 1024
    print(f'  Size: {size_kb} KB')


if __name__ == '__main__':
    main()
