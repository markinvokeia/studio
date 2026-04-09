"""
generate_agent_plan.py
Genera el documento Word "Agente IA — Plan de Implementación" para Invoke IA.
Uso: python3 scripts/generate_agent_plan.py
Output: docs/agente-ia-plan-implementacion.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "..", "docs", "agente-ia-plan-implementacion.docx")

# ── Paleta de colores ──────────────────────────────────────────────────────────
COLOR_PRIMARY   = RGBColor(0x1E, 0x40, 0xAF)   # azul oscuro
COLOR_ACCENT    = RGBColor(0x06, 0xB6, 0xD4)   # cyan
COLOR_DARK      = RGBColor(0x1E, 0x29, 0x3B)   # casi negro
COLOR_TIP_BG    = RGBColor(0xEF, 0xF6, 0xFF)   # azul muy suave
COLOR_TIP_TEXT  = RGBColor(0x1E, 0x40, 0xAF)
COLOR_WARN_TEXT = RGBColor(0x92, 0x40, 0x0E)
COLOR_GRAY      = RGBColor(0x6B, 0x72, 0x80)
COLOR_GREEN     = RGBColor(0x06, 0x6B, 0x3B)
COLOR_GREEN_BG  = RGBColor(0xF0, 0xFD, 0xF4)
COLOR_PURPLE    = RGBColor(0x6D, 0x28, 0xD9)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def remove_cell_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_tip(doc, text: str, kind: str = "tip"):
    table = doc.add_table(rows=1, cols=1)
    remove_cell_borders(table)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    if kind == "tip":
        set_cell_bg(cell, "EFF6FF")
        icon = "💡 Consejo: "
        color = COLOR_TIP_TEXT
    elif kind == "warn":
        set_cell_bg(cell, "FFF7ED")
        icon = "⚠️  Nota: "
        color = COLOR_WARN_TEXT
    else:
        set_cell_bg(cell, "F0FDF4")
        icon = "✅ Requisito: "
        color = COLOR_GREEN
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(icon)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(9.5)
    run2 = p.add_run(text)
    run2.font.color.rgb = color
    run2.font.size = Pt(9.5)
    doc.add_paragraph()


def add_section_heading(doc, text: str, level: int = 2):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.color.rgb = COLOR_PRIMARY


def add_chapter_heading(doc, number: int, title: str, subtitle: str = ""):
    doc.add_page_break()
    p = doc.add_heading(f"{number}. {title}", level=1)
    for run in p.runs:
        run.font.color.rgb = COLOR_PRIMARY
    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.runs[0].font.color.rgb = COLOR_GRAY
        sub.runs[0].font.italic = True
        sub.runs[0].font.size = Pt(10)
    doc.add_paragraph()


def add_bullet(doc, text: str, level: int = 1, bold_prefix: str = ""):
    style = "List Bullet" if level == 1 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(10.5)
    p.add_run(text).font.size = Pt(10.5)


def add_code_block(doc, text: str):
    """Agrega un bloque de código con fondo gris."""
    table = doc.add_table(rows=1, cols=1)
    remove_cell_borders(table)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, "1E293B")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    doc.add_paragraph()


def add_use_case_card(doc, cu_id: str, title: str, tier: str, trigger_examples: list,
                      apis: list, tool_name: str, confirmation: bool, notes: str = ""):
    """Genera una tarjeta estructurada para un caso de uso."""
    add_section_heading(doc, f"{cu_id}: {title}")

    # Tier badge
    tier_colors = {
        "TIER 1": ("FF0000", "Impacto Inmediato"),
        "TIER 2": ("FF6B00", "Flujos Completos"),
        "TIER 3": ("0066CC", "Atención al Paciente"),
        "TIER 4": ("6D28D9", "Inteligencia y Reportes"),
    }
    tier_hex, tier_label = tier_colors.get(tier, ("666666", tier))
    p_tier = doc.add_paragraph()
    r = p_tier.add_run(f"  {tier} — {tier_label}  ")
    r.font.size = Pt(8.5)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Ejemplos de activación
    p_ex = doc.add_paragraph()
    r2 = p_ex.add_run("Frases de activación: ")
    r2.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = COLOR_DARK
    r3 = p_ex.add_run('"' + '" · "'.join(trigger_examples) + '"')
    r3.font.size = Pt(10)
    r3.font.italic = True
    r3.font.color.rgb = COLOR_GRAY

    # Tool responsable
    p_tool = doc.add_paragraph()
    rt = p_tool.add_run("Tool / Sub-workflow responsable: ")
    rt.bold = True
    rt.font.size = Pt(10)
    p_tool.add_run(tool_name).font.size = Pt(10)

    # Confirmación requerida
    p_conf = doc.add_paragraph()
    rc = p_conf.add_run("Confirmación requerida: ")
    rc.bold = True
    rc.font.size = Pt(10)
    conf_text = "SÍ — el agente presenta un resumen antes de ejecutar" if confirmation else "NO — solo lectura, se ejecuta directo"
    p_conf.add_run(conf_text).font.size = Pt(10)

    # APIs
    p_apis = doc.add_paragraph()
    r_apis = p_apis.add_run("Cadena de APIs:")
    r_apis.bold = True
    r_apis.font.size = Pt(10)
    for api in apis:
        sub = doc.add_paragraph(style="List Bullet 2")
        sub.paragraph_format.space_before = Pt(1)
        sub.paragraph_format.space_after = Pt(1)
        sub.add_run(api).font.size = Pt(9.5)

    if notes:
        add_tip(doc, notes, kind="tip")

    doc.add_paragraph()


def build_document():
    doc = Document()

    # ── Configuración de página ────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ── PORTADA ────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run("INVOKE IA")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Agente IA — Plan de Implementación")
    run2.font.size = Pt(20)
    run2.font.color.rgb = COLOR_DARK

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Casos de uso, sub-workflows, APIs y cambios al orquestador Main")
    run3.font.size = Pt(12)
    run3.font.color.rgb = COLOR_GRAY
    run3.font.italic = True

    doc.add_paragraph("\n" * 4)
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Versión 1.0  ·  Abril 2026")
    run4.font.size = Pt(10)
    run4.font.color.rgb = COLOR_GRAY

    # ══════════════════════════════════════════════════════════════════════════
    # CAPÍTULO 1 — ARQUITECTURA ACTUAL DEL ORQUESTADOR MAIN
    # ══════════════════════════════════════════════════════════════════════════
    add_chapter_heading(doc, 1, "Arquitectura Actual del Orquestador Main",
                        "Estado del flujo principal de n8n al inicio del proyecto")

    doc.add_paragraph(
        "El orquestador Main es el flujo central de n8n que recibe mensajes desde múltiples "
        "canales, normaliza la entrada, valida la identidad del usuario y delega la ejecución "
        "a un Agente LLM (Manager) equipado con herramientas especializadas."
    )

    add_section_heading(doc, "1.1  Canales de Entrada")
    channels = [
        ("WhatsApp Oficial", "WhatsApp Trigger (n8n nativo)", "texto, audio, imagen"),
        ("WhatsApp Evolution API", "Webhook POST /api/chat/message (JWT auth)", "texto, audio"),
        ("Telegram", "Telegram Trigger", "texto, audio, foto"),
        ("Gmail", "Gmail Trigger", "email completo"),
        ("Web Chat Widget", "Webhook chat (ID: a8ad846b)", "texto"),
    ]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Canal", "Trigger", "Tipos de media"]):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_bg(hdr[i], "1E40AF")
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for name, trigger, media in channels:
        row = tbl.add_row().cells
        row[0].text = name
        row[1].text = trigger
        row[2].text = media
    doc.add_paragraph()

    add_section_heading(doc, "1.2  Enrutamiento y Procesamiento de Media")
    doc.add_paragraph(
        "El nodo 'Redirect Message Types1' es un Switch con 10 salidas que separa los mensajes "
        "según canal y tipo de contenido. Los mensajes de audio pasan por un pipeline de "
        "transcripción (Gemini Audio) y los de imagen por un pipeline de descripción visual "
        "(Gemini Vision). Todos convergen en el nodo 'Request' con el siguiente payload:"
    )
    add_code_block(doc,
        "{\n"
        '  "message":     "<texto normalizado>",\n'
        '  "source":      "whatsapp | telegram | gmail | chat | evolution",\n'
        '  "sessionId":   "<identificador único de sesión>",\n'
        '  "chatInputId": "<id para Web Chat>"\n'
        "}"
    )

    add_section_heading(doc, "1.3  Capa de Seguridad")
    steps_sec = [
        ("Get_User_Data", "Sub-workflow gcnS82LQE4Mocaek — busca al usuario por número de teléfono/email en Invoke IA y devuelve su token JWT y clinic_id."),
        ("Switch Is Registered", "Si el usuario NO está registrado → rama 'User Onboarding' (sub-workflow s6ANusf96nLvxFSb). Si SÍ está registrado → continúa al Log Access."),
        ("Log Access", "INSERT en tabla PostgreSQL 'access_logs' con sessionId, userId, source, timestamp. El registro se actualiza al final (Update Access) o en caso de error (Update Error Access)."),
    ]
    for title_s, desc in steps_sec:
        p_s = doc.add_paragraph(style="List Number")
        r_s = p_s.add_run(title_s)
        r_s.bold = True
        r_s.font.size = Pt(10.5)
        p_s.add_run(f" — {desc}").font.size = Pt(10.5)

    add_section_heading(doc, "1.4  El Manager (Agente LLM)")
    doc.add_paragraph(
        "El Manager es un nodo LangChain Agent (@n8n/n8n-nodes-langchain.agent) configurado con:"
    )
    add_bullet(doc, "Modelo: Google Gemini 2.5 Flash")
    add_bullet(doc, "Memoria: Postgres Chat Memory — últimos 20 mensajes, clave = sessionId")
    add_bullet(doc, "Herramientas registradas: 16 toolWorkflow nodes (ver sección 1.5)")
    add_bullet(doc, "System prompt: define identidad, idioma, reglas de confirmación y lista de capacidades")
    add_tip(doc,
        "El Manager NO ejecuta código directamente. Interpreta la intención del usuario y "
        "llama a la herramienta apropiada pasando un parámetro 'query' en lenguaje natural. "
        "Cada herramienta es un sub-workflow n8n que interpreta ese query, llama las APIs de "
        "Invoke IA y devuelve el resultado formateado.")

    add_section_heading(doc, "1.5  Herramientas Registradas (estado actual)")
    current_tools = [
        ("calendarAgent", "Gestión de Google Calendar (crear/editar/cancelar eventos)", "Existente"),
        ("emailAgent", "Envío y consulta de emails", "Existente"),
        ("researchAgent", "Búsqueda web y síntesis de información", "Existente"),
        ("notesAgent", "Notas internas y recordatorios del staff", "Existente"),
        ("reminderAgent", "Recordatorios programados", "Existente"),
        ("makeOutboundCall", "Llamadas salientes", "Existente"),
        ("getContactsTool", "Búsqueda de pacientes/contactos en Invoke IA", "Existente"),
        ("getServicesTool", "Catálogo de servicios", "Existente"),
        ("usersAdminTool", "Administración de usuarios del sistema", "Existente"),
        ("quoteBillingAndPaymentTool", "Presupuestos, facturas y pagos", "Existente"),
        ("anamnesisAgent", "Historia clínica y anamnesis", "Existente"),
        ("getPadecimientosTool", "Catálogo de padecimientos", "Existente"),
        ("getMedicinesTool", "Catálogo de medicamentos", "Existente"),
        ("getHorariosAtencionInformacionContacto", "Horarios y datos de contacto (SQL directo)", "Existente"),
        ("sendTextMessagesToUsers", "Envío de mensajes a pacientes", "Existente"),
        ("clinicSessionAndNoteTakerAgent", "Sesiones clínicas y notas de atención", "Existente"),
    ]
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    hdr2 = tbl2.rows[0].cells
    for i, h in enumerate(["Herramienta", "Descripción actual", "Estado"]):
        hdr2[i].text = h
        for p in hdr2[i].paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_bg(hdr2[i], "1E40AF")
        for para in hdr2[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for name, desc, status in current_tools:
        row = tbl2.add_row().cells
        row[0].text = name
        row[1].text = desc
        row[2].text = status
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # CAPÍTULO 2 — NUEVOS AGENTES Y MCPs A CREAR
    # ══════════════════════════════════════════════════════════════════════════
    add_chapter_heading(doc, 2, "Nuevos Agentes y MCPs a Crear",
                        "Sub-workflows especializados que el Manager delegará según la intención")

    doc.add_paragraph(
        "Para cubrir todos los casos de uso identificados, se crean 6 nuevos sub-workflows "
        "(registrados como tools en Main) agrupados por dominio funcional. Cada uno recibe un "
        "parámetro 'query' con la intención en lenguaje natural, el token JWT del usuario y el "
        "clinic_id, y devuelve una respuesta ya formateada para el canal."
    )

    # ── appointmentsAgent ──────────────────────────────────────────────────────
    add_section_heading(doc, "2.1  appointmentsAgent  (NUEVO)")
    doc.add_paragraph(
        "Dominio: gestión completa de la agenda médica. Cubre CU-01, CU-03 y CU-05."
    )
    add_bullet(doc, "Consultar agenda del día o de un rango de fechas")
    add_bullet(doc, "Buscar disponibilidad de un doctor")
    add_bullet(doc, "Crear una nueva cita (con validación de conflictos)")
    add_bullet(doc, "Cancelar o reprogramar una cita existente")
    add_bullet(doc, "Marcar una cita/atención como completada")
    doc.add_paragraph()
    p_io = doc.add_paragraph()
    r_io = p_io.add_run("Entradas requeridas:")
    r_io.bold = True
    add_bullet(doc, "query: texto en lenguaje natural con la intención", level=2)
    add_bullet(doc, "token: JWT del usuario autenticado", level=2)
    add_bullet(doc, "clinic_id: ID de la clínica", level=2)
    add_bullet(doc, "today: fecha actual (YYYY-MM-DD) para resolver referencias relativas", level=2)

    doc.add_paragraph()
    p_apis2 = doc.add_paragraph()
    r_apis2 = p_apis2.add_run("APIs internas que consume:")
    r_apis2.bold = True
    for api in [
        "GET /users_appointments — listar citas con filtros de fecha/doctor/paciente",
        "GET /filter_users?search={nombre} — resolver nombre a user_id",
        "GET /users/doctors — listar médicos disponibles",
        "GET /services?search={nombre} — resolver nombre de servicio",
        "GET /appointments_availability — verificar huecos libres",
        "POST /appointments/upsert — crear o actualizar cita",
        "POST /appointments/delete — cancelar cita",
        "GET /calendars — obtener calendarios y sus colores",
    ]:
        add_bullet(doc, api, level=2)

    add_tip(doc,
        "Para CU-03 (crear cita), si el horario solicitado no está disponible, el sub-workflow "
        "debe llamar a /appointments_availability y ofrecer los 3 próximos slots libres antes "
        "de responder al usuario.", kind="tip")

    # ── financialQueryAgent ──────────────────────────────────────────────────
    add_section_heading(doc, "2.2  financialQueryAgent  (NUEVO)")
    doc.add_paragraph(
        "Dominio: consultas financieras de pacientes. Cubre CU-02. Solo lectura — no modifica datos."
    )
    add_bullet(doc, "Deuda total y facturas pendientes de un paciente")
    add_bullet(doc, "Saldo a favor / prepagos disponibles")
    add_bullet(doc, "Historial de pagos realizados")
    add_bullet(doc, "Estado de un presupuesto u orden específica")

    doc.add_paragraph()
    p_fa = doc.add_paragraph()
    r_fa = p_fa.add_run("APIs internas que consume:")
    r_fa.bold = True
    for api in [
        "GET /filter_users?search={nombre} — resolver nombre a user_id",
        "GET /user_financial?user_id={id} — resumen financiero completo",
        "GET /user_invoices?user_id={id} — facturas con estado",
        "GET /user_payments?user_id={id} — historial de pagos",
        "GET /user_credit?user_id={id} — saldo a favor y créditos",
        "GET /quotes?user_id={id} — presupuestos del paciente",
        "GET /all_orders?user_id={id} — órdenes vinculadas",
    ]:
        add_bullet(doc, api, level=2)

    # ── paymentRegistrationTool ──────────────────────────────────────────────
    add_section_heading(doc, "2.3  paymentRegistrationTool  (NUEVO)")
    doc.add_paragraph(
        "Dominio: registro de pagos desde conversación. Cubre CU-04. "
        "REQUIERE confirmación explícita antes de ejecutar escritura."
    )
    add_bullet(doc, "Registrar pago parcial o total sobre una factura")
    add_bullet(doc, "Aplicar prepago (saldo disponible) a factura existente")
    add_bullet(doc, "Registrar abono sin factura asociada (prepago)")

    doc.add_paragraph()
    p_pr = doc.add_paragraph()
    r_pr = p_pr.add_run("Flujo de confirmación obligatorio:")
    r_pr.bold = True
    r_pr.font.color.rgb = COLOR_WARN_TEXT
    add_code_block(doc,
        "1. Agente recibe: 'Registra pago $2500 UYU efectivo para López'\n"
        "2. Sub-workflow resuelve: user_id, factura pendiente, método de pago\n"
        "3. Sub-workflow devuelve DRAFT al Manager:\n"
        "   '¿Confirmas pago de $2.500 UYU en Efectivo para FAC-0042 (Juan López)?'\n"
        "4. Usuario responde SI → sub-workflow ejecuta POST /invoice/payment\n"
        "5. Sub-workflow devuelve confirmación: 'Pago registrado. Saldo: $2.000 UYU'"
    )

    doc.add_paragraph()
    p_pra = doc.add_paragraph()
    r_pra = p_pra.add_run("APIs internas que consume:")
    r_pra.bold = True
    for api in [
        "GET /filter_users?search={nombre} — user_id",
        "GET /user_invoices?user_id={id}&status=pending — facturas pendientes",
        "GET /metodospago/all — métodos de pago disponibles",
        "POST /invoice/payment — registrar pago en factura",
        "GET /user_financial?user_id={id} — saldo final post-pago",
    ]:
        add_bullet(doc, api, level=2)

    # ── cashierAgent ──────────────────────────────────────────────────────────
    add_section_heading(doc, "2.4  cashierAgent  (NUEVO)")
    doc.add_paragraph(
        "Dominio: gestión de caja y sesiones de caja. Cubre CU-08."
    )
    add_bullet(doc, "Consultar si la caja está abierta y su saldo actual")
    add_bullet(doc, "Abrir sesión de caja con fondo inicial")
    add_bullet(doc, "Guiar el proceso de cierre paso a paso")
    add_bullet(doc, "Consultar movimientos del día")

    doc.add_paragraph()
    p_ca = doc.add_paragraph()
    r_ca = p_ca.add_run("APIs internas que consume:")
    r_ca.bold = True
    for api in [
        "GET /cash-session/active — sesión de caja activa",
        "GET /cash_points/search — puntos de caja registrados",
        "GET /cash-session/movements — movimientos de la sesión",
        "GET /cash-session/prefill — datos prefill para cierre",
        "POST /cash-session/open — abrir sesión con fondo inicial",
        "POST /cash-session/declare — declarar efectivo físico contado",
        "POST /cash-session/close — cerrar la sesión",
        "GET /cotizaciones — tasas de cambio UYU/USD del día",
    ]:
        add_bullet(doc, api, level=2)

    # ── dashboardAgent ────────────────────────────────────────────────────────
    add_section_heading(doc, "2.5  dashboardAgent  (NUEVO)")
    doc.add_paragraph(
        "Dominio: KPIs y estadísticas del negocio. Cubre CU-06 y CU-15. Solo lectura."
    )
    add_bullet(doc, "Resumen financiero del período (ingresos, crecimiento)")
    add_bullet(doc, "Métricas de citas (show rate, cancelaciones)")
    add_bullet(doc, "Tasa de conversión presupuesto → venta")
    add_bullet(doc, "Servicios más demandados")
    add_bullet(doc, "Pacientes nuevos vs recurrentes")
    add_bullet(doc, "Estado de facturas (pagadas / pendientes / vencidas)")
    add_bullet(doc, "Alertas pendientes del sistema")

    doc.add_paragraph()
    p_da = doc.add_paragraph()
    r_da = p_da.add_run("APIs internas que consume:")
    r_da.bold = True
    for api in [
        "GET /dashboard_summary — KPIs principales del período",
        "GET /dashboard_sales_summary — tendencia de ingresos",
        "GET /dashboard_sales_by_service — ventas por servicio/tratamiento",
        "GET /dashboard_invoice_status — estado de facturación",
        "GET /dashboard_new_vs_recurring_patients — nuevos vs recurrentes",
        "GET /users_appointments (today) — citas del día para briefing matutino",
        "GET /cash-session/active — estado de caja para briefing",
        "GET /system/alert-instances?status=PENDING — alertas sin resolver",
    ]:
        add_bullet(doc, api, level=2)

    # ── patientPortalAgent ────────────────────────────────────────────────────
    add_section_heading(doc, "2.6  patientPortalAgent  (NUEVO — canal externo)")
    doc.add_paragraph(
        "Dominio: autoservicio para PACIENTES externos via WhatsApp. Cubre CU-11, CU-12, CU-13. "
        "Opera con token de servicio restringido. Los pacientes solo pueden ver sus propios datos "
        "(identificados por número de teléfono)."
    )
    add_bullet(doc, "Consultar próxima cita propia")
    add_bullet(doc, "Solicitar nuevo turno (queda en estado 'pendiente' para confirmar por staff)")
    add_bullet(doc, "Consultar precio de servicios del catálogo")
    add_bullet(doc, "Cancelar propia cita (con anticipación mínima configurable)")

    doc.add_paragraph()
    p_ppa = doc.add_paragraph()
    r_ppa = p_ppa.add_run("APIs internas que consume:")
    r_ppa.bold = True
    for api in [
        "GET /filter_users?phone={telefono} — resolver teléfono a user_id (identificación)",
        "GET /users_appointments?patient_id={id} — citas propias",
        "GET /appointments_availability — disponibilidad para auto-agendado",
        "POST /appointments/upsert (status=pending) — solicitud de turno",
        "GET /services — catálogo de precios",
    ]:
        add_bullet(doc, api, level=2)

    add_tip(doc,
        "Los pacientes externos NUNCA pueden: ver datos de otros pacientes, registrar pagos, "
        "acceder a historia clínica de otros ni modificar información del sistema. El token de "
        "servicio del patientPortalAgent debe tener permisos mínimos de solo lectura.", kind="warn")

    # ══════════════════════════════════════════════════════════════════════════
    # CAPÍTULO 3 — CASOS DE USO DETALLADOS
    # ══════════════════════════════════════════════════════════════════════════
    add_chapter_heading(doc, 3, "Casos de Uso Detallados",
                        "Especificación completa de cada caso de uso con APIs, entidades y flujo")

    # ── TIER 1 ──────────────────────────────────────────────────────────────
    doc.add_heading("TIER 1 — Impacto Inmediato (alta frecuencia diaria)", level=2)
    for run in doc.paragraphs[-1].runs:
        run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    add_use_case_card(
        doc,
        cu_id="CU-01",
        title="Consulta de Agenda del Día",
        tier="TIER 1",
        trigger_examples=["¿Qué citas hay hoy?", "¿A qué hora tiene cita María?", "Agenda del Dr. García para mañana"],
        apis=[
            "GET /users_appointments?startingDateAndTime=...&endingDateAndTime=...&calendar_ids=...",
            "GET /filter_users?search={nombre} — solo si se filtra por paciente",
            "GET /calendars — para resolver nombres y colores de calendarios",
        ],
        tool_name="appointmentsAgent (NUEVO)",
        confirmation=False,
        notes="Respuesta formateada: lista de citas con hora · paciente · doctor · estado. "
              "Si no hay citas, responder 'No hay citas programadas para ese período.'"
    )

    add_use_case_card(
        doc,
        cu_id="CU-02",
        title="Consulta de Deuda / Estado Financiero de Paciente",
        tier="TIER 1",
        trigger_examples=["¿Cuánto debe Pedro Ruiz?", "Facturas pendientes de López", "Estado de cuenta de Ana García"],
        apis=[
            "GET /filter_users?search={nombre} → obtiene user_id",
            "GET /user_financial?user_id={id} → deuda total, saldo, facturado",
            "GET /user_invoices?user_id={id} → facturas con estado y montos",
            "GET /user_credit?user_id={id} → créditos y saldo a favor",
        ],
        tool_name="financialQueryAgent (NUEVO)",
        confirmation=False,
        notes="Si hay múltiples pacientes con el mismo nombre, listar opciones y pedir al staff que confirme cuál."
    )

    add_use_case_card(
        doc,
        cu_id="CU-03",
        title="Crear Cita por Conversación",
        tier="TIER 1",
        trigger_examples=["Agenda a Juan con el Dr. García mañana a las 10 para limpieza", "Nuevo turno para Ana Pérez el viernes"],
        apis=[
            "GET /filter_users?search={paciente} → patient_id",
            "GET /users/doctors → lista de médicos",
            "GET /services?search={servicio} → service_id",
            "GET /appointments_availability?doctor_id=...&date=... → slots libres",
            "POST /appointments/upsert → crear cita confirmada",
        ],
        tool_name="appointmentsAgent (NUEVO)",
        confirmation=True,
        notes="Si el horario solicitado no está libre, ofrecer los 3 próximos slots disponibles. "
              "Siempre confirmar antes de crear: 'Voy a crear cita para [paciente] con [doctor] el [fecha] a las [hora] — ¿Confirmas?'"
    )

    add_use_case_card(
        doc,
        cu_id="CU-04",
        title="Registrar Pago desde Chat",
        tier="TIER 1",
        trigger_examples=["Registra pago de $2500 UYU efectivo para López", "Cobro de la factura de Ana García"],
        apis=[
            "GET /filter_users?search={nombre} → user_id",
            "GET /user_invoices?user_id={id}&status=pending → facturas pendientes",
            "GET /metodospago/all → métodos de pago disponibles",
            "POST /invoice/payment → registrar pago",
            "GET /user_financial?user_id={id} → saldo final post-pago",
        ],
        tool_name="paymentRegistrationTool (NUEVO)",
        confirmation=True,
        notes="El agente SIEMPRE presenta un resumen antes de ejecutar el pago. "
              "El sub-workflow retorna un objeto DRAFT que el Manager presenta al usuario para confirmación."
    )

    add_use_case_card(
        doc,
        cu_id="CU-05",
        title="Completar Atención (Flujo Multi-Paso)",
        tier="TIER 1",
        trigger_examples=["Completa la atención de Juan Pérez, realizó extracción pieza 26", "Marcar como completado al paciente de las 10"],
        apis=[
            "GET /users_appointments?date=today&status=confirmed → cita activa del paciente",
            "POST /sesiones/upsert → crear sesión clínica con procedimiento y notas",
            "POST /appointments/upsert (status=completed) → cerrar la cita",
            "GET /order_items?appointment_id={id} → ítems de orden vinculados",
            "POST /orders/upsert (item_status=completed) → marcar ítem como completado",
        ],
        tool_name="clinicSessionAndNoteTakerAgent (EXTENDER existente)",
        confirmation=True,
        notes="Este flujo extiende el sub-workflow existente 'clinicSessionAndNoteTakerAgent'. "
              "Agregar lógica para cerrar la cita y marcar ítems de orden como completados."
    )

    add_use_case_card(
        doc,
        cu_id="CU-06",
        title="Resumen del Día (Briefing Matutino)",
        tier="TIER 1",
        trigger_examples=["Dame el resumen del día", "¿Cómo estamos hoy?", "Briefing matutino"],
        apis=[
            "GET /users_appointments?date=today → citas del día",
            "GET /cash-session/active → estado de caja",
            "GET /system/alert-instances?status=PENDING&priority=HIGH → alertas sin resolver",
            "GET /dashboard_summary → KPIs del período",
        ],
        tool_name="dashboardAgent (NUEVO)",
        confirmation=False,
        notes="Se puede disparar automáticamente al inicio del día vía cron trigger en n8n. "
              "Respuesta estructurada: citas del día · estado de caja · alertas críticas · primer paciente."
    )

    # ── TIER 2 ──────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("TIER 2 — Flujos Completos Automatizados", level=2)
    for run in doc.paragraphs[-1].runs:
        run.font.color.rgb = RGBColor(0xEA, 0x58, 0x0C)

    add_use_case_card(
        doc,
        cu_id="CU-07",
        title="Crear Presupuesto por Conversación",
        tier="TIER 2",
        trigger_examples=["Crea un presupuesto para Carlos con ortodoncia y extracción", "Nuevo presupuesto en USD para Ana García"],
        apis=[
            "GET /filter_users?search={paciente} → user_id",
            "GET /services — catálogo de servicios con precios",
            "POST /quotes/upsert → crear presupuesto con cabecera",
            "POST /quote/lines/upsert → agregar ítems al presupuesto",
            "GET /config/currencies → monedas disponibles (UYU/USD)",
        ],
        tool_name="quoteBillingAndPaymentTool (EXTENDER existente)",
        confirmation=True,
        notes="El sub-workflow existente ya maneja presupuestos. Extender para soportar creación "
              "conversacional completa incluyendo múltiples ítems en un solo mensaje."
    )

    add_use_case_card(
        doc,
        cu_id="CU-08",
        title="Apertura / Cierre de Caja Asistido",
        tier="TIER 2",
        trigger_examples=["Abrir caja con $10.000 UYU", "Cerrar caja", "¿Está abierta la caja?"],
        apis=[
            "GET /cash-session/active → verificar si ya hay sesión",
            "GET /cash_points/search → punto de caja a usar",
            "GET /cotizaciones → tasa de cambio del día",
            "POST /cash-session/open → abrir con fondo inicial",
            "GET /cash-session/prefill → datos para asistir el cierre",
            "GET /cash-session/movements → movimientos a revisar",
            "POST /cash-session/declare → registrar conteo físico",
            "POST /cash-session/close → cerrar sesión",
        ],
        tool_name="cashierAgent (NUEVO)",
        confirmation=True,
        notes="El cierre es guiado paso a paso: el agente pregunta el efectivo contado por denominación, "
              "calcula diferencias y solicita justificación si el descuadre supera el umbral."
    )

    add_use_case_card(
        doc,
        cu_id="CU-09",
        title="Pipeline Presupuesto → Orden → Factura",
        tier="TIER 2",
        trigger_examples=["El paciente aceptó el presupuesto PRES-0023", "Convierte la cotización de López a factura"],
        apis=[
            "GET /quotes?id={id} → detalle del presupuesto",
            "POST /quote/confirm → confirmar presupuesto (genera orden automáticamente)",
            "GET /all_orders?quote_id={id} → orden generada",
            "POST /invoices/upsert → crear factura desde orden con ítems completados",
        ],
        tool_name="quoteBillingAndPaymentTool (EXTENDER existente)",
        confirmation=True,
        notes="Verificar si la API /quote/confirm genera la orden automáticamente o si "
              "se requiere un POST /orders/upsert separado. Adaptar el flujo según el comportamiento real."
    )

    add_use_case_card(
        doc,
        cu_id="CU-10",
        title="Alertas Proactivas por WhatsApp",
        tier="TIER 2",
        trigger_examples=["(trigger automático cada hora)", "¿Hay alertas pendientes?"],
        apis=[
            "GET /system/alert-instances?status=PENDING&priority=HIGH → alertas críticas",
            "GET /filter_users?role=admin → ID del administrador receptor",
            "POST /system/alert-instances/send-whatsapp → marcar alerta como notificada",
        ],
        tool_name="Nuevo cron trigger en n8n → Manager directo",
        confirmation=False,
        notes="Este CU no pasa por el Manager como herramienta. Es un flujo n8n separado con "
              "cron trigger que consulta alertas y envía mensajes push al admin via WhatsApp/Evolution."
    )

    # ── TIER 3 ──────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("TIER 3 — Atención al Paciente (Canal Externo)", level=2)
    for run in doc.paragraphs[-1].runs:
        run.font.color.rgb = RGBColor(0x15, 0x63, 0xEB)

    add_use_case_card(
        doc,
        cu_id="CU-11",
        title="Paciente Consulta su Próxima Cita",
        tier="TIER 3",
        trigger_examples=["¿Cuándo es mi próxima cita?", "Hola, quiero saber mi turno"],
        apis=[
            "GET /filter_users?phone={telefono} → identificar paciente por número WA",
            "GET /users_appointments?patient_id={id}&status=confirmed&from=today → próxima cita",
        ],
        tool_name="patientPortalAgent (NUEVO)",
        confirmation=False,
        notes="La identificación del paciente se hace por el número de teléfono del mensaje entrante. "
              "Si el número no está registrado, responder con info de contacto de la clínica."
    )

    add_use_case_card(
        doc,
        cu_id="CU-12",
        title="Paciente Solicita Nuevo Turno",
        tier="TIER 3",
        trigger_examples=["Quiero una cita para limpieza la próxima semana", "¿Tienen disponibilidad para el viernes?"],
        apis=[
            "GET /filter_users?phone={telefono} → patient_id",
            "GET /services?search={servicio} → service_id",
            "GET /appointments_availability → 3 slots disponibles más próximos",
            "POST /appointments/upsert (status=pending) → solicitud pendiente de confirmación por staff",
        ],
        tool_name="patientPortalAgent (NUEVO)",
        confirmation=True,
        notes="La cita se crea con status='pending'. El staff debe confirmarla desde el sistema. "
              "Una notificación se envía automáticamente a recepción al crear la solicitud."
    )

    add_use_case_card(
        doc,
        cu_id="CU-13",
        title="Paciente Consulta Precio de Servicios",
        tier="TIER 3",
        trigger_examples=["¿Cuánto cuesta un implante?", "Precios de ortodoncia", "¿Tienen blanqueamiento?"],
        apis=[
            "GET /services?search={servicio} → precio y descripción del catálogo",
            "GET /user_quotes?user_id={id} → presupuesto personalizado si existe (opcional)",
        ],
        tool_name="patientPortalAgent (NUEVO)",
        confirmation=False,
        notes="Si el paciente ya tiene un presupuesto personalizado, mencionar que puede consultar "
              "su propuesta específica. Los precios del catálogo son referenciales."
    )

    add_use_case_card(
        doc,
        cu_id="CU-14",
        title="Recordatorio Automático de Cita (Proactivo)",
        tier="TIER 3",
        trigger_examples=["(cron: 24h y 2h antes de cada cita)", "(trigger automático)"],
        apis=[
            "GET /users_appointments?from=tomorrow&to=tomorrow → citas del día siguiente",
            "GET /filter_users?user_id={id} → email y teléfono del paciente",
            "POST /system/alert-instances/send-whatsapp → enviar recordatorio",
        ],
        tool_name="Nuevo cron trigger en n8n (independiente del Manager)",
        confirmation=False,
        notes="Implementar como flujo n8n separado con dos triggers: 24 horas antes y 2 horas antes. "
              "Incluir botones de respuesta: 'Confirmar' / 'Cancelar'. Si el paciente cancela, "
              "actualizar estado de la cita y notificar a recepción."
    )

    # ── TIER 4 ──────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("TIER 4 — Inteligencia y Reportes", level=2)
    for run in doc.paragraphs[-1].runs:
        run.font.color.rgb = COLOR_PURPLE

    add_use_case_card(
        doc,
        cu_id="CU-15",
        title="Dashboard Conversacional",
        tier="TIER 4",
        trigger_examples=["¿Cómo van las ventas este mes?", "KPIs del mes", "¿Cuántos pacientes nuevos tuvimos?"],
        apis=[
            "GET /dashboard_summary?period={mes} → KPIs principales",
            "GET /dashboard_sales_summary → tendencia de ingresos",
            "GET /dashboard_sales_by_service → servicios más demandados",
            "GET /dashboard_invoice_status → estado de facturación",
            "GET /dashboard_new_vs_recurring_patients → nuevos vs recurrentes",
        ],
        tool_name="dashboardAgent (NUEVO)",
        confirmation=False,
        notes="El LLM reformatea los datos numéricos en lenguaje natural con contexto y comparaciones. "
              "Ejemplo: 'Este mes tuvieron $48.500 UYU, un 12% más que el mes pasado.'"
    )

    add_use_case_card(
        doc,
        cu_id="CU-16",
        title="Consulta de Historia Clínica para el Doctor",
        tier="TIER 4",
        trigger_examples=["¿Qué antecedentes tiene Juan Pérez?", "Alergias de Ana García", "Historial de tratamientos de López"],
        apis=[
            "GET /filter_users?search={nombre} → user_id",
            "GET /clinic-history/{user_id} → historia clínica completa",
            "GET /antecedentes_personales?user_id={id} → antecedentes personales",
            "GET /antecedentes_alergias?user_id={id} → alergias registradas",
            "GET /antecedentes_medicamentos?user_id={id} → medicación actual",
            "GET /patient_sessions?user_id={id} → historial de sesiones/tratamientos",
        ],
        tool_name="anamnesisAgent (EXTENDER existente)",
        confirmation=False,
        notes="El sub-workflow existente 'anamnesisAgent' ya gestiona historia clínica. Extender "
              "para soportar consultas de síntesis rápida: el LLM resume la información más relevante."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # CAPÍTULO 4 — CAMBIOS AL ORQUESTADOR MAIN
    # ══════════════════════════════════════════════════════════════════════════
    add_chapter_heading(doc, 4, "Cambios Requeridos al Orquestador Main",
                        "Nuevos nodos toolWorkflow y actualizaciones al system prompt")

    add_section_heading(doc, "4.1  Nuevos Nodos toolWorkflow a Agregar")
    doc.add_paragraph(
        "Cada nuevo sub-workflow se registra como un nodo toolWorkflow en el Manager de Main. "
        "El patrón es idéntico a los tools existentes: el Manager llama a cada tool con un "
        "parámetro 'query' en lenguaje natural y el tool devuelve el resultado."
    )

    new_tools = [
        ("appointmentsAgent", "Crear, consultar, cancelar y completar citas. Verificar disponibilidad. Ver agenda del día.", "CU-01, CU-03, CU-05"),
        ("financialQueryAgent", "Consultar deuda, facturas pendientes, saldo y créditos de un paciente.", "CU-02"),
        ("paymentRegistrationTool", "Registrar pagos sobre facturas. SIEMPRE requiere confirmación antes de ejecutar.", "CU-04"),
        ("cashierAgent", "Abrir, consultar y cerrar sesión de caja. Guiar el proceso de arqueo.", "CU-08"),
        ("dashboardAgent", "Resumen del día, KPIs, estadísticas de ventas y métricas clínicas.", "CU-06, CU-15"),
        ("patientPortalAgent", "Autoservicio para pacientes externos: consultar cita propia, solicitar turno, ver precios.", "CU-11, CU-12, CU-13"),
    ]

    tbl3 = doc.add_table(rows=1, cols=3)
    tbl3.style = "Table Grid"
    hdr3 = tbl3.rows[0].cells
    for i, h in enumerate(["Nombre del Tool", "Descripción para el Manager", "Casos de Uso"]):
        hdr3[i].text = h
        for p in hdr3[i].paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_bg(hdr3[i], "1E40AF")
        for para in hdr3[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for name, desc, cus in new_tools:
        row = tbl3.add_row().cells
        row[0].text = name
        row[1].text = desc
        row[2].text = cus
    doc.add_paragraph()

    add_section_heading(doc, "4.2  Sub-workflows Existentes a Extender")
    extend_tools = [
        ("clinicSessionAndNoteTakerAgent", "Agregar lógica para cerrar la cita asociada y marcar ítems de orden como completados al registrar una sesión.", "CU-05"),
        ("quoteBillingAndPaymentTool", "Agregar soporte para: (a) creación conversacional de presupuestos multi-ítem, (b) pipeline confirm→orden→factura en un solo flujo.", "CU-07, CU-09"),
        ("anamnesisAgent", "Agregar modo de consulta rápida: sintetizar los antecedentes más relevantes de un paciente en pocas líneas.", "CU-16"),
    ]
    tbl4 = doc.add_table(rows=1, cols=3)
    tbl4.style = "Table Grid"
    hdr4 = tbl4.rows[0].cells
    for i, h in enumerate(["Sub-workflow", "Extensión requerida", "Casos de Uso"]):
        hdr4[i].text = h
        for p in hdr4[i].paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_bg(hdr4[i], "1E40AF")
        for para in hdr4[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for name, ext, cus in extend_tools:
        row = tbl4.add_row().cells
        row[0].text = name
        row[1].text = ext
        row[2].text = cus
    doc.add_paragraph()

    add_section_heading(doc, "4.3  Flujos Independientes (sin pasar por el Manager)")
    doc.add_paragraph("Algunos casos de uso son proactivos (inician sin mensaje del usuario) y se implementan como flujos n8n separados:")
    independent = [
        ("Recordatorios de Cita (CU-14)", "Cron cada 1 hora. Busca citas en próximas 24h y 2h. Envía WhatsApp con botones Confirmar/Cancelar."),
        ("Alertas Proactivas (CU-10)", "Cron cada 1 hora. Busca alertas CRITICAL/HIGH sin resolver. Envía resumen al administrador."),
        ("Briefing Automático (CU-06 variante)", "Cron 08:00 AM. Compone resumen del día y lo envía al group/chat del staff."),
    ]
    for name, desc in independent:
        p_ind = doc.add_paragraph(style="List Bullet")
        r_ind = p_ind.add_run(name + ": ")
        r_ind.bold = True
        r_ind.font.size = Pt(10.5)
        p_ind.add_run(desc).font.size = Pt(10.5)

    add_section_heading(doc, "4.4  Actualización del System Prompt del Manager")
    doc.add_paragraph("El system prompt actual debe actualizarse para incluir las nuevas herramientas. Fragmento a agregar:")
    add_code_block(doc,
        "## Nuevas Capacidades\n\n"
        "- appointmentsAgent: usa para TODO lo relacionado con agenda, citas y disponibilidad.\n"
        "  Ejemplos: ver agenda del día, crear cita, cancelar, marcar como completada.\n\n"
        "- financialQueryAgent: usa para consultas financieras de pacientes (deuda, facturas,\n"
        "  saldo a favor). SOLO LECTURA, no ejecuta pagos.\n\n"
        "- paymentRegistrationTool: usa ÚNICAMENTE para registrar pagos. SIEMPRE presenta\n"
        "  resumen al usuario y espera confirmación explícita (SI/NO) antes de ejecutar.\n\n"
        "- cashierAgent: usa para abrir, cerrar o consultar la sesión de caja.\n\n"
        "- dashboardAgent: usa para KPIs, estadísticas, resumen del día y métricas.\n\n"
        "- patientPortalAgent: usa cuando el source es 'whatsapp' o 'evolution' y el usuario\n"
        "  NO es staff (identificado como paciente externo por teléfono).\n\n"
        "## Reglas de Confirmación\n"
        "Antes de ejecutar cualquier acción que modifique datos (crear cita, registrar pago,\n"
        "abrir/cerrar caja, crear presupuesto), SIEMPRE presenta un resumen y pide confirmación."
    )

    # ══════════════════════════════════════════════════════════════════════════
    # CAPÍTULO 5 — TABLA DE INTEGRACIÓN COMPLETA
    # ══════════════════════════════════════════════════════════════════════════
    add_chapter_heading(doc, 5, "Tabla de Integración Completa",
                        "Mapeo de intención → tool → APIs → confirmación requerida")

    rows_data = [
        ("CU-01", "consultar_agenda", "appointmentsAgent", "/users_appointments + /calendars", "No"),
        ("CU-02", "consultar_deuda_paciente", "financialQueryAgent", "/filter_users + /user_financial + /user_invoices", "No"),
        ("CU-03", "crear_cita", "appointmentsAgent", "/filter_users + /users/doctors + /services + /appointments_availability + /appointments/upsert", "Sí"),
        ("CU-04", "registrar_pago", "paymentRegistrationTool", "/filter_users + /user_invoices + /metodospago/all + /invoice/payment", "Sí"),
        ("CU-05", "completar_atencion", "clinicSessionAndNoteTakerAgent", "/users_appointments + /sesiones/upsert + /appointments/upsert + /orders/upsert", "Sí"),
        ("CU-06", "resumen_del_dia", "dashboardAgent", "/users_appointments + /cash-session/active + /system/alert-instances + /dashboard_summary", "No"),
        ("CU-07", "crear_presupuesto", "quoteBillingAndPaymentTool", "/filter_users + /services + /quotes/upsert + /quote/lines/upsert", "Sí"),
        ("CU-08", "abrir_caja / cerrar_caja", "cashierAgent", "/cash-session/active + /cash_points/search + /cash-session/open|close", "Sí"),
        ("CU-09", "confirmar_presupuesto", "quoteBillingAndPaymentTool", "/quote/confirm + /all_orders + /invoices/upsert", "Sí"),
        ("CU-10", "alerta_proactiva", "Flujo cron independiente", "/system/alert-instances + /send-whatsapp", "No"),
        ("CU-11", "paciente_consulta_cita", "patientPortalAgent", "/filter_users?phone + /users_appointments", "No"),
        ("CU-12", "paciente_solicita_turno", "patientPortalAgent", "/filter_users?phone + /appointments_availability + /appointments/upsert(pending)", "Sí"),
        ("CU-13", "consultar_precio_servicio", "patientPortalAgent", "/services?search", "No"),
        ("CU-14", "recordatorio_cita", "Flujo cron independiente", "/users_appointments + /filter_users + /send-whatsapp", "No"),
        ("CU-15", "consultar_estadisticas", "dashboardAgent", "/dashboard_summary + /dashboard_sales_* + /dashboard_invoice_status", "No"),
        ("CU-16", "consultar_historial_clinico", "anamnesisAgent", "/filter_users + /clinic-history + /patient_sessions", "No"),
    ]

    tbl5 = doc.add_table(rows=1, cols=5)
    tbl5.style = "Table Grid"
    hdr5 = tbl5.rows[0].cells
    for i, h in enumerate(["CU", "Intención", "Tool/Sub-workflow", "APIs principales", "Confirmación"]):
        hdr5[i].text = h
        for p in hdr5[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_bg(hdr5[i], "1E40AF")
        for para in hdr5[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for cu, intent, tool, apis, conf in rows_data:
        row = tbl5.add_row().cells
        row[0].text = cu
        row[1].text = intent
        row[2].text = tool
        row[3].text = apis
        row[4].text = conf
        for cell in row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8.5)
        if conf == "Sí":
            set_cell_bg(row[4], "FEF9C3")
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # CAPÍTULO 6 — SEGURIDAD Y CONSIDERACIONES
    # ══════════════════════════════════════════════════════════════════════════
    add_chapter_heading(doc, 6, "Seguridad y Consideraciones de Implementación",
                        "Principios de seguridad, manejo de errores y rollout sugerido")

    add_section_heading(doc, "6.1  Modelo de Tokens")
    doc.add_paragraph(
        "El agente opera con dos tipos de token según el canal y el tipo de usuario:"
    )
    add_bullet(doc, "Token de staff (personal): JWT del usuario autenticado en Invoke IA. "
               "Permisos completos según el rol del usuario (doctor, administrativo, gerente).")
    add_bullet(doc, "Token de servicio (paciente externo): JWT de un usuario 'Agente IA' "
               "con permisos mínimos de solo lectura para datos propios del paciente. "
               "Se usa exclusivamente en el patientPortalAgent.")
    add_tip(doc,
        "Crear en Invoke IA un usuario con rol 'Agente Paciente' y permisos limitados: "
        "ver_citas_propias, ver_servicios, solicitar_turno. "
        "Nunca usar el token de administrador para el canal de pacientes externos.", kind="warn")

    add_section_heading(doc, "6.2  Patrón de Confirmación para Acciones Destructivas")
    doc.add_paragraph("Todo sub-workflow que modifique datos debe implementar el siguiente patrón de dos pasos:")
    add_code_block(doc,
        "PASO 1 — Preparar (siempre):\n"
        "  • El sub-workflow resuelve entidades (user_id, doctor_id, etc.)\n"
        "  • Retorna al Manager un objeto DRAFT con el resumen de la acción\n"
        "  • El Manager presenta el resumen al usuario y espera respuesta\n\n"
        "PASO 2 — Ejecutar (solo con SI/Confirmar explícito):\n"
        "  • El Manager llama al mismo sub-workflow con flag execute=true\n"
        "  • El sub-workflow ejecuta la escritura en la API\n"
        "  • Retorna confirmación con número de documento generado\n\n"
        "ALTERNATIVA: implementar con nodo 'Wait for Webhook' de n8n para el paso 2."
    )

    add_section_heading(doc, "6.3  Manejo de Ambigüedad")
    add_bullet(doc, "Múltiples pacientes con el mismo nombre: listar las opciones con identificador, pedir confirmación de cuál.")
    add_bullet(doc, "Fecha/hora ambigua ('mañana', 'la próxima semana'): resolver usando el campo 'today' del contexto y pedir confirmación.")
    add_bullet(doc, "Intención desconocida: si el Manager no puede clasificar, responder con opciones concretas de lo que puede hacer.")
    add_bullet(doc, "Fallo de API: capturar errores en el sub-workflow y devolver mensaje amigable al usuario.")

    add_section_heading(doc, "6.4  Plan de Rollout Sugerido")
    phases = [
        ("Semana 1-2", "Crear appointmentsAgent y financialQueryAgent. Agregar al Main. Probar CU-01 y CU-02."),
        ("Semana 3", "Crear dashboardAgent. Probar CU-06 y CU-15. Configurar briefing matutino cron."),
        ("Semana 4", "Crear paymentRegistrationTool. Probar CU-04 con confirmación. Revisar patrón de dos pasos."),
        ("Semana 5-6", "Extender clinicSessionAndNoteTakerAgent y quoteBillingAndPaymentTool para CU-05, CU-07, CU-09."),
        ("Semana 7", "Crear cashierAgent para CU-08. Probar apertura y cierre guiado."),
        ("Semana 8-9", "Crear patientPortalAgent para CU-11, CU-12, CU-13. Implementar flujo cron CU-14."),
        ("Semana 10", "Implementar alertas proactivas CU-10. Ajustar system prompt del Manager. QA completo."),
    ]
    for phase, desc in phases:
        p_ph = doc.add_paragraph(style="List Number")
        r_ph = p_ph.add_run(phase + ": ")
        r_ph.bold = True
        r_ph.font.size = Pt(10.5)
        p_ph.add_run(desc).font.size = Pt(10.5)

    add_tip(doc,
        "Cada nuevo tool debe registrarse en el Manager con una descripción clara de CUÁNDO usarlo. "
        "El Manager usa esas descripciones para decidir qué tool llamar — una descripción imprecisa "
        "causará que el Manager llame la herramienta equivocada.", kind="warn")

    # ── Guardar ──────────────────────────────────────────────────────────────
    output_dir = os.path.dirname(OUTPUT_PATH)
    os.makedirs(output_dir, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"✅ Documento generado: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
