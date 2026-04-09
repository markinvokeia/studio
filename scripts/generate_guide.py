"""
generate_guide.py
Genera el documento Word "Guía de Uso — Invoke IA" para nuevos usuarios.
Uso: python3 scripts/generate_guide.py
Output: docs/guia-uso-invoke-ia.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "..", "docs", "guia-uso-invoke-ia.docx")

# ── Paleta de colores ──────────────────────────────────────────────────────────
COLOR_PRIMARY   = RGBColor(0x1E, 0x40, 0xAF)   # azul oscuro
COLOR_ACCENT    = RGBColor(0x06, 0xB6, 0xD4)   # cyan
COLOR_DARK      = RGBColor(0x1E, 0x29, 0x3B)   # casi negro
COLOR_TIP_BG    = RGBColor(0xEF, 0xF6, 0xFF)   # azul muy suave
COLOR_TIP_TEXT  = RGBColor(0x1E, 0x40, 0xAF)
COLOR_WARN_TEXT = RGBColor(0x92, 0x40, 0x0E)
COLOR_GRAY      = RGBColor(0x6B, 0x72, 0x80)


def set_cell_bg(cell, hex_color: str):
    """Pinta el fondo de una celda de tabla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def remove_cell_borders(table):
    """Quita todos los bordes de una tabla (para cajas de tip/aviso)."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_tip(doc, text: str, kind: str = "tip"):
    """Agrega un recuadro de tip o advertencia."""
    table = doc.add_table(rows=1, cols=1)
    remove_cell_borders(table)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    if kind == "tip":
        set_cell_bg(cell, "EFF6FF")
        icon = "💡 Consejo: "
        color = COLOR_TIP_TEXT
    else:
        set_cell_bg(cell, "FFF7ED")
        icon = "⚠️  Nota: "
        color = COLOR_WARN_TEXT
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(icon)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(9.5)
    run2 = p.add_run(text)
    run2.font.color.rgb = color
    run2.font.size = Pt(9.5)
    doc.add_paragraph()  # espacio después del recuadro


def add_step(doc, number: int, title: str, details: list[str]):
    """Agrega un paso numerado con sub-bullets."""
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    for detail in details:
        sub = doc.add_paragraph(style="List Bullet 2")
        sub.paragraph_format.space_before = Pt(1)
        sub.paragraph_format.space_after = Pt(1)
        sub.add_run(detail).font.size = Pt(10)


def add_section_heading(doc, text: str):
    """Agrega un encabezado de sección (H2)."""
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.color.rgb = COLOR_PRIMARY


def add_module_heading(doc, number: int, title: str, subtitle: str = ""):
    """Agrega un encabezado de módulo (H1) con banda de color."""
    doc.add_page_break()
    p = doc.add_heading(f"Módulo {number} — {title}", level=1)
    for run in p.runs:
        run.font.color.rgb = COLOR_PRIMARY
    if subtitle:
        sub = doc.add_paragraph(subtitle)
        sub.runs[0].font.color.rgb = COLOR_GRAY
        sub.runs[0].font.italic = True
        sub.runs[0].font.size = Pt(10)
    doc.add_paragraph()


def build_document():
    doc = Document()

    # ── Configuración de página ────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # ── Estilos base ───────────────────────────────────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # ── PORTADA ────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    run = p.add_run("INVOKE IA")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Guía de Uso para Nuevos Usuarios")
    run2.font.size = Pt(20)
    run2.font.color.rgb = COLOR_DARK

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("Gestión integral de clínicas odontológicas y de salud")
    run3.font.size = Pt(12)
    run3.font.color.rgb = COLOR_GRAY
    run3.font.italic = True

    doc.add_paragraph("\n" * 4)
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Versión 1.0  ·  Abril 2026")
    run4.font.size = Pt(10)
    run4.font.color.rgb = COLOR_GRAY

    # ── INTRODUCCIÓN ──────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Introducción", level=1)
    intro = doc.add_paragraph(
        "Invoke IA es una plataforma de gestión integral diseñada para clínicas odontológicas "
        "y de salud. Permite administrar pacientes, citas, presupuestos, facturación, caja y "
        "mucho más desde un único lugar.\n\n"
        "Esta guía está organizada en el orden recomendado para comenzar a usar el sistema: "
        "primero las configuraciones que deben realizarse una sola vez, luego las operaciones "
        "del día a día."
    )
    intro.paragraph_format.space_after = Pt(8)

    doc.add_heading("Cómo usar esta guía", level=2)
    bullets = [
        "Sigue los módulos en orden la primera vez que configures el sistema.",
        "Cada flujo incluye pasos numerados y consejos prácticos.",
        "Los módulos 2 al 5 son los de uso diario; los módulos 6–9 son opcionales según tu clínica.",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        p.paragraph_format.space_after = Pt(3)

    # ══════════════════════════════════════════════════════════════════════════
    # MÓDULO 1 — CONFIGURACIÓN INICIAL
    # ══════════════════════════════════════════════════════════════════════════
    add_module_heading(doc, 1, "Configuración Inicial",
                       "Realiza estos pasos una sola vez antes de comenzar a operar.")

    # 1.1 Perfil de la clínica
    add_section_heading(doc, "1.1  Perfil de la Clínica")
    doc.add_paragraph(
        "Define los datos básicos de tu clínica: nombre, dirección, teléfono, email y logo. "
        "Esta información aparecerá en presupuestos, facturas y comunicaciones al paciente."
    )
    add_step(doc, 1, "Ir a Configuración → Clínica", [
        "En el menú lateral, haz clic en Configuración.",
        "Selecciona la opción Clínica.",
    ])
    add_step(doc, 2, "Completar los datos", [
        "Ingresa nombre, dirección, teléfono, email y sitio web.",
        "Sube el logo de la clínica (recomendado: PNG con fondo transparente).",
    ])
    add_step(doc, 3, "Guardar", ["Haz clic en Guardar cambios."])
    add_tip(doc, "El logo debe tener al menos 200×200 px para verse nítido en los documentos.")

    # 1.2 Monedas
    add_section_heading(doc, "1.2  Monedas y Tipos de Cambio")
    doc.add_paragraph(
        "El sistema soporta múltiples monedas (ej. USD y UYU). Configura las monedas que "
        "usará tu clínica y define el tipo de cambio vigente."
    )
    add_step(doc, 1, "Ir a Configuración → Monedas", [
        "En el menú lateral, ve a Configuración → Monedas.",
    ])
    add_step(doc, 2, "Agregar o activar monedas", [
        "Haz clic en Nueva moneda o activa las existentes.",
        "Define el símbolo (USD, UYU, etc.) y el tipo de cambio frente a la moneda base.",
    ])
    add_step(doc, 3, "Actualizar el tipo de cambio periódicamente", [
        "Ingresa el valor actualizado y guarda para que los presupuestos reflejen el cambio correcto.",
    ])
    add_tip(doc, "Actualiza el tipo de cambio al inicio de cada semana si trabajas con tarifas en USD.")

    # 1.3 Médicos
    add_section_heading(doc, "1.3  Médicos y Dentistas")
    doc.add_paragraph(
        "Registra a cada profesional que atiende en la clínica. Los médicos se usan para asignar "
        "citas, filtrar la agenda y asociar servicios."
    )
    add_step(doc, 1, "Ir a Configuración → Médicos", [])
    add_step(doc, 2, "Crear un nuevo médico", [
        "Haz clic en Nuevo médico.",
        "Completa: nombre, apellido, especialidad, teléfono y email.",
        "Asigna un color de identificación para la agenda (cada médico tendrá su color en el calendario).",
    ])
    add_step(doc, 3, "Guardar", ["Haz clic en Guardar."])
    add_tip(doc, "Asigna colores distintos a cada médico para identificarlos fácilmente en la agenda.")

    # 1.4 Disponibilidad
    add_section_heading(doc, "1.4  Disponibilidad de Médicos")
    doc.add_paragraph(
        "Define los días y horarios en que cada médico puede recibir pacientes. "
        "El sistema usará esta información para validar que las citas se agenden dentro "
        "del horario disponible."
    )
    add_step(doc, 1, "Ir a Configuración → Disponibilidad de médicos", [])
    add_step(doc, 2, "Seleccionar el médico", [
        "Elige el médico al que quieres configurar la disponibilidad.",
    ])
    add_step(doc, 3, "Definir horarios por día", [
        "Para cada día de la semana, marca si trabaja ese día.",
        "Ingresa hora de inicio y fin de atención.",
        "Opcionalmente, define un intervalo de descanso (almuerzo, etc.).",
    ])
    add_step(doc, 4, "Guardar los cambios", [])
    add_tip(doc, "Usa Excepciones de disponibilidad para registrar vacaciones o días especiales de un médico.",
            kind="note")

    # 1.5 Calendarios y Horarios
    add_section_heading(doc, "1.5  Calendarios y Horarios de la Clínica")
    doc.add_paragraph(
        "Configura el horario general de apertura de la clínica y los calendarios que se "
        "usarán en la agenda de citas."
    )
    add_step(doc, 1, "Ir a Configuración → Horarios", [
        "Define el horario de atención de la clínica por día de la semana.",
    ])
    add_step(doc, 2, "Crear un Calendario", [
        "Ve a Configuración → Calendarios.",
        "Crea un calendario (ej. 'Consultorio Principal') y asócialo al horario configurado.",
    ])
    add_step(doc, 3, "Asociar médicos al calendario", [
        "Cada médico puede estar asociado a uno o más calendarios.",
    ])

    # 1.6 Feriados
    add_section_heading(doc, "1.6  Feriados y Días No Laborables")
    doc.add_paragraph(
        "Registra los feriados del año para que el sistema no permita agendar citas esos días."
    )
    add_step(doc, 1, "Ir a Configuración → Feriados", [])
    add_step(doc, 2, "Agregar un feriado", [
        "Haz clic en Nuevo feriado.",
        "Ingresa la fecha y una descripción (ej. 'Año Nuevo').",
        "Haz clic en Guardar.",
    ])
    add_tip(doc, "Carga todos los feriados del año al inicio para evitar sorpresas en la agenda.")

    # 1.7 Catálogo de Servicios
    add_section_heading(doc, "1.7  Catálogo de Servicios")
    doc.add_paragraph(
        "Define todos los procedimientos o servicios que ofrece tu clínica, con su precio. "
        "Estos servicios se usarán al crear presupuestos y al agendar citas."
    )
    add_step(doc, 1, "Ir a Ventas → Servicios", [])
    add_step(doc, 2, "Crear un servicio", [
        "Haz clic en Nuevo servicio.",
        "Ingresa: nombre, descripción, precio y moneda.",
        "Opcionalmente, asigna una duración estimada (para bloquear tiempo en la agenda).",
        "Haz clic en Guardar.",
    ])
    add_tip(doc, "Puedes crear servicios en múltiples monedas. El sistema convertirá automáticamente según el tipo de cambio configurado.")

    # 1.8 Métodos de Pago
    add_section_heading(doc, "1.8  Métodos de Pago")
    doc.add_paragraph(
        "Configura los métodos de pago que acepta la clínica (efectivo, tarjeta, transferencia, etc.)."
    )
    add_step(doc, 1, "Ir a Ventas → Métodos de pago", [])
    add_step(doc, 2, "Activar o crear métodos", [
        "Activa los métodos disponibles en el sistema o crea nuevos según necesidad.",
        "Define el nombre y si genera recibo imprimible.",
    ])

    # 1.9 Cajas Físicas
    add_section_heading(doc, "1.9  Cajas Físicas")
    doc.add_paragraph(
        "Si la clínica tiene más de una caja o punto de cobro, regístralos aquí. "
        "Cada cajero deberá seleccionar su caja al iniciar sesión."
    )
    add_step(doc, 1, "Ir a Caja → Puntos de venta", [])
    add_step(doc, 2, "Crear una caja", [
        "Haz clic en Nueva caja.",
        "Ingresa el nombre (ej. 'Caja Principal', 'Recepción 2').",
        "Haz clic en Guardar.",
    ])

    # 1.10 Secuencias
    add_section_heading(doc, "1.10  Secuencias de Documentos")
    doc.add_paragraph(
        "Define el prefijo y número inicial para la numeración automática de presupuestos, "
        "órdenes y facturas."
    )
    add_step(doc, 1, "Ir a Configuración → Secuencias", [])
    add_step(doc, 2, "Configurar cada tipo de documento", [
        "Para cada tipo (Presupuesto, Orden, Factura), define el prefijo (ej. 'PRES-', 'FAC-') "
        "y el número de inicio (ej. 1000).",
        "Haz clic en Guardar.",
    ])
    add_tip(doc, "Configura esto antes de crear el primer presupuesto para mantener una numeración ordenada desde el inicio.")

    # 1.11 Usuarios y Roles
    add_section_heading(doc, "1.11  Usuarios del Sistema y Roles")
    doc.add_paragraph(
        "Crea cuentas para cada persona que usará el sistema y asígnales un rol con los permisos "
        "correspondientes (admin, recepcionista, cajero, médico, etc.)."
    )
    add_step(doc, 1, "Crear roles", [
        "Ve a Sistema → Roles.",
        "Crea los roles necesarios (ej. 'Recepcionista', 'Cajero') y selecciona los permisos de cada uno.",
        "Haz clic en Guardar.",
    ])
    add_step(doc, 2, "Crear usuarios", [
        "Ve a Sistema → Usuarios.",
        "Haz clic en Nuevo usuario.",
        "Ingresa nombre, email, teléfono y asigna un rol.",
        "El usuario recibirá un email para establecer su contraseña.",
    ])
    add_tip(doc, "Asigna el rol de mínimos privilegios necesarios: un recepcionista no necesita acceso a configuración del sistema.", kind="note")

    # ══════════════════════════════════════════════════════════════════════════
    # MÓDULO 2 — PACIENTES
    # ══════════════════════════════════════════════════════════════════════════
    add_module_heading(doc, 2, "Gestión de Pacientes",
                       "Crea y administra la ficha de cada paciente.")

    add_section_heading(doc, "2.1  Crear un Nuevo Paciente")
    doc.add_paragraph(
        "Antes de agendar una cita o emitir un presupuesto, el paciente debe estar registrado en el sistema."
    )
    add_step(doc, 1, "Ir a Pacientes", ["En el menú lateral, haz clic en Pacientes."])
    add_step(doc, 2, "Nuevo paciente", ["Haz clic en el botón Nuevo paciente (ícono +)."])
    add_step(doc, 3, "Completar el formulario", [
        "Datos obligatorios: nombre, apellido y al menos un teléfono o email.",
        "Datos opcionales: fecha de nacimiento, documento de identidad, dirección, notas.",
    ])
    add_step(doc, 4, "Guardar", ["Haz clic en Guardar. El paciente aparecerá en la lista."])
    add_tip(doc, "Busca primero antes de crear para evitar duplicados: usa el buscador en la parte superior de la lista.")

    add_section_heading(doc, "2.2  Ver y Editar la Ficha del Paciente")
    doc.add_paragraph(
        "La ficha del paciente centraliza toda su información: datos personales, historial "
        "financiero, citas, presupuestos, facturas y pagos."
    )
    add_step(doc, 1, "Seleccionar un paciente", [
        "En la lista de Pacientes, haz clic sobre el nombre del paciente.",
        "Se abrirá el panel lateral con su ficha.",
    ])
    add_step(doc, 2, "Navegar por las pestañas", [
        "Resumen financiero: balance total, deuda pendiente, monto facturado.",
        "Citas: historial y próximas citas del paciente.",
        "Presupuestos: todos los presupuestos emitidos.",
        "Facturas y Pagos: documentos de cobro e historial de pagos.",
        "Historia Clínica: sesiones, diagnósticos y tratamientos.",
    ])
    add_step(doc, 3, "Editar datos", [
        "Haz clic en el ícono de editar (lápiz) en el encabezado.",
        "Modifica los campos necesarios y guarda.",
    ])

    add_section_heading(doc, "2.3  Historia Clínica")
    doc.add_paragraph(
        "Registra cada visita clínica del paciente: diagnósticos, tratamientos realizados, "
        "medicamentos indicados y archivos adjuntos."
    )
    add_step(doc, 1, "Abrir la historia clínica", [
        "Desde la ficha del paciente, ve a la pestaña Historia Clínica.",
        "O ve a Historia Clínica en el menú y busca al paciente.",
    ])
    add_step(doc, 2, "Nueva sesión clínica", [
        "Haz clic en Nueva sesión.",
        "Ingresa la fecha de la visita y el médico tratante.",
        "Describe los motivos de consulta y observaciones.",
    ])
    add_step(doc, 3, "Agregar diagnósticos y tratamientos", [
        "Selecciona las afecciones del catálogo (o créalas si no existen).",
        "Registra los tratamientos realizados.",
    ])
    add_step(doc, 4, "Recetar medicamentos", [
        "En la sección Medicamentos, selecciona del catálogo y define dosis e indicaciones.",
    ])
    add_step(doc, 5, "Adjuntar archivos", [
        "Adjunta radiografías, análisis u otros documentos relevantes.",
        "Haz clic en Guardar sesión.",
    ])
    add_tip(doc, "Configura el Catálogo Clínico (afecciones, medicamentos, condiciones dentales) en el menú Catálogo Clínico antes de registrar sesiones.", kind="note")

    # ══════════════════════════════════════════════════════════════════════════
    # MÓDULO 3 — AGENDA Y CITAS
    # ══════════════════════════════════════════════════════════════════════════
    add_module_heading(doc, 3, "Agenda y Citas",
                       "Gestiona el calendario de atención de la clínica.")

    add_section_heading(doc, "3.1  Crear una Cita")
    add_step(doc, 1, "Ir a Citas", ["En el menú lateral, haz clic en Citas."])
    add_step(doc, 2, "Seleccionar una franja horaria", [
        "En la vista de calendario, haz clic sobre el horario deseado en la columna del médico.",
        "O haz clic en el botón Nueva cita.",
    ])
    add_step(doc, 3, "Completar el formulario", [
        "Paciente: busca y selecciona al paciente (o créalo en el momento).",
        "Médico: selecciona el profesional que atenderá.",
        "Servicio: elige el tratamiento o procedimiento.",
        "Fecha y hora: confirma el horario seleccionado.",
        "Notas opcionales para el médico.",
    ])
    add_step(doc, 4, "Confirmar", ["Haz clic en Guardar. La cita aparecerá en el calendario con el color del médico."])
    add_tip(doc, "Si el paciente no tiene horario disponible con ese médico, el sistema te avisará con un aviso de conflicto.", kind="note")

    add_section_heading(doc, "3.2  Gestionar el Calendario")
    add_step(doc, 1, "Cambiar la vista", [
        "Usa los botones Día / Semana / Mes para cambiar la vista.",
        "La vista Semana es la más útil para ver la disponibilidad de todos los médicos.",
    ])
    add_step(doc, 2, "Filtrar por médico", [
        "Usa el panel de filtros para mostrar solo los calendarios de los médicos que te interesen.",
    ])
    add_step(doc, 3, "Reprogramar una cita", [
        "Arrastra y suelta la cita a la nueva franja horaria (drag & drop).",
        "O haz clic en la cita y edita la fecha/hora.",
    ])
    add_step(doc, 4, "Cancelar una cita", [
        "Haz clic en la cita → Editar → cambia el estado a Cancelada → Guardar.",
    ])
    add_tip(doc, "Las citas canceladas permanecen en el historial del paciente para referencia.")

    add_section_heading(doc, "3.3  Enlazar una Cita con un Presupuesto")
    doc.add_paragraph(
        "Al crear o editar una cita es posible asociarla a un presupuesto existente o crear uno nuevo "
        "en el mismo flujo, sin salir del formulario de cita."
    )
    add_step(doc, 1, "Abrir el formulario de cita", [
        "Desde el Calendario, haz clic en un hueco horario o sobre una cita existente.",
        "O usa el botón PROGRAMAR desde el ítem de servicio en una Orden de Venta.",
    ])
    add_step(doc, 2, "Ir a la sección 'Presupuesto asociado'", [
        "Dentro del formulario de cita, busca el campo Presupuesto asociado.",
    ])
    add_step(doc, 3, "Seleccionar presupuesto existente", [
        "Despliega la lista y elige un presupuesto activo del paciente.",
        "El sistema filtra automáticamente los presupuestos del paciente seleccionado.",
    ])
    add_step(doc, 4, "O crear un presupuesto nuevo", [
        "Haz clic en '+ Nuevo presupuesto' para abrir el diálogo rápido de creación.",
        "Agrega los ítems, moneda y notas. Al confirmar, queda vinculado automáticamente.",
    ])
    add_step(doc, 5, "Guardar la cita", [
        "Haz clic en Guardar. El presupuesto vinculado quedará visible en el detalle de la cita.",
    ])
    add_tip(doc, "Vincular la cita al presupuesto permite rastrear qué atenciones corresponden a cada propuesta comercial y facilita la facturación posterior.")

    add_section_heading(doc, "3.4  Completar una Cita con Sesión Clínica")
    doc.add_paragraph(
        "Cuando se realiza la atención al paciente, la cita puede vincularse a una Sesión Clínica "
        "que documenta lo ejecutado. Esto cierra el ciclo: de la cita agendada a la intervención registrada."
    )
    add_step(doc, 1, "Desde el Calendario", [
        "Localiza la cita ejecutada y ábrela.",
        "Haz clic en 'Registrar sesión' dentro del formulario de cita.",
        "Ingresa los procedimientos, observaciones y adjuntos relevantes.",
        "Al confirmar, la sesión queda vinculada y la cita pasa a estado Completada.",
    ])
    add_step(doc, 2, "Desde el botón COMPLETAR (en servicios de una Orden)", [
        "Ve a la Orden de Venta y localiza el ítem de servicio atendido.",
        "Haz clic en el botón COMPLETAR del ítem.",
        "Se abre el diálogo de sesión clínica. Completa los datos de la atención.",
        "Al confirmar: la sesión se registra en la historia del paciente y el ítem queda marcado como Completado.",
    ])
    add_tip(doc, "Solo los ítems en estado Completado son elegibles para incluirse en una factura. Completa el ítem antes de facturar.", kind="note")

    # ══════════════════════════════════════════════════════════════════════════
    # MÓDULO 4 — FLUJO COMERCIAL
    # ══════════════════════════════════════════════════════════════════════════
    add_module_heading(doc, 4, "Flujo Comercial: Presupuesto → Factura → Cobro",
                       "El ciclo completo de ventas de la clínica.")

    doc.add_paragraph(
        "El flujo comercial estándar sigue estas etapas:"
    )
    flow_table = doc.add_table(rows=1, cols=5)
    flow_table.style = "Table Grid"
    stages = ["Presupuesto", "→", "Orden", "→", "Factura + Pago"]
    for i, cell in enumerate(flow_table.rows[0].cells):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(stages[i])
        run.bold = (i % 2 == 0)
        run.font.size = Pt(10)
        if i % 2 == 0:
            set_cell_bg(cell, "DBEAFE")
    doc.add_paragraph()

    add_section_heading(doc, "4.1  Crear un Presupuesto")
    doc.add_paragraph(
        "El presupuesto es una propuesta de tratamiento con precios que se presenta al paciente "
        "antes de iniciar el trabajo."
    )
    add_step(doc, 1, "Ir a Ventas → Presupuestos", [])
    add_step(doc, 2, "Nuevo presupuesto", ["Haz clic en Nuevo presupuesto."])
    add_step(doc, 3, "Seleccionar paciente", [
        "Busca y selecciona al paciente.",
        "Define la moneda del presupuesto.",
    ])
    add_step(doc, 4, "Agregar servicios", [
        "Haz clic en Agregar ítem.",
        "Selecciona el servicio del catálogo (el precio se carga automáticamente).",
        "Ajusta la cantidad o el precio si es necesario.",
        "Repite para todos los tratamientos incluidos.",
    ])
    add_step(doc, 5, "Revisar y enviar", [
        "Revisa el total.",
        "Cambia el estado a Enviado para notificar al paciente por email.",
        "Haz clic en Guardar.",
    ])
    add_tip(doc, "Puedes guardar el presupuesto en estado Borrador y enviarlo más tarde cuando esté listo.")

    add_section_heading(doc, "4.2  Confirmar un Presupuesto y Crear una Orden")
    doc.add_paragraph(
        "Cuando el paciente acepta el presupuesto, se convierte en una Orden de trabajo."
    )
    add_step(doc, 1, "Abrir el presupuesto aceptado", [
        "En Ventas → Presupuestos, busca el presupuesto.",
        "Haz clic en él para abrirlo.",
    ])
    add_step(doc, 2, "Cambiar estado a Aceptado", [
        "Cambia el estado del presupuesto a Aceptado.",
        "Haz clic en Convertir a Orden.",
    ])
    add_step(doc, 3, "Confirmar la orden", [
        "Se creará automáticamente una orden con los mismos ítems.",
        "Revisa y haz clic en Guardar.",
    ])
    add_tip(doc, "Las órdenes quedan en Ventas → Órdenes. Desde ahí se generará la factura.")

    add_section_heading(doc, "4.3  Generar una Factura")
    add_step(doc, 1, "Ir a Ventas → Órdenes", [])
    add_step(doc, 2, "Abrir la orden correspondiente", ["Haz clic en la orden."])
    add_step(doc, 3, "Convertir a Factura", [
        "Haz clic en Convertir a Factura.",
        "Revisa los ítems y el total.",
        "Haz clic en Guardar. La factura quedará en estado Pendiente de pago.",
    ])
    add_tip(doc, "Desde la factura puedes hacer clic en Imprimir para generar el PDF del documento.")

    add_section_heading(doc, "4.4  Registrar un Pago")
    add_step(doc, 1, "Abrir la factura", ["Ve a Ventas → Facturas y abre la factura pendiente."])
    add_step(doc, 2, "Registrar pago", [
        "Haz clic en Registrar pago.",
        "Ingresa el monto, la fecha y el método de pago (efectivo, tarjeta, etc.).",
        "Si el pago es parcial, ingresa el monto pagado; el sistema calculará el saldo pendiente.",
    ])
    add_step(doc, 3, "Guardar", [
        "Haz clic en Guardar. La factura cambiará a estado Pagada (o Pago parcial si aplica).",
    ])
    add_tip(doc, "Los pagos también pueden registrarse desde la sección Ventas → Pagos o directamente desde la ficha del paciente.")

    # ══════════════════════════════════════════════════════════════════════════
    # MÓDULO 5 — CAJA
    # ══════════════════════════════════════════════════════════════════════════
    add_module_heading(doc, 5, "Gestión de Caja",
                       "Control del flujo de efectivo en la clínica.")

    add_section_heading(doc, "5.1  Abrir una Sesión de Caja")
    doc.add_paragraph(
        "Cada turno de trabajo en caja comienza con una apertura que registra el saldo inicial en efectivo."
    )
    add_step(doc, 1, "Ir a Caja", ["En el menú lateral, haz clic en Caja."])
    add_step(doc, 2, "Iniciar sesión", [
        "Haz clic en Abrir sesión.",
        "Selecciona la caja física (si hay más de una).",
        "Ingresa el saldo inicial en efectivo (el dinero con que comienzas el turno).",
        "Haz clic en Confirmar apertura.",
    ])
    add_tip(doc, "Solo puede haber una sesión activa por caja a la vez.")

    add_section_heading(doc, "5.2  Registrar Transacciones Misceláneas")
    doc.add_paragraph(
        "Para ingresos o egresos que no provienen de una factura de venta "
        "(ej. compra de insumos en efectivo, gastos varios)."
    )
    add_step(doc, 1, "Ir a Caja → Transacciones misceláneas", [])
    add_step(doc, 2, "Nueva transacción", [
        "Haz clic en Nueva transacción.",
        "Selecciona tipo: Ingreso o Egreso.",
        "Elige la categoría (configurable en Caja → Categorías).",
        "Ingresa el monto, descripción y método de pago.",
        "Haz clic en Guardar.",
    ])

    add_section_heading(doc, "5.3  Cerrar la Sesión de Caja")
    add_step(doc, 1, "Ir a Caja", [])
    add_step(doc, 2, "Cerrar sesión", [
        "Haz clic en Cerrar sesión.",
        "Ingresa el saldo final contado en efectivo.",
        "El sistema calculará diferencias respecto al saldo esperado.",
        "Haz clic en Confirmar cierre.",
    ])
    add_step(doc, 3, "Revisar el resumen", [
        "El sistema mostrará el resumen: ingresos, egresos, saldo esperado vs. real.",
        "Puedes imprimir el cierre de caja.",
    ])
    add_tip(doc, "Revisa el historial de sesiones en Caja → Sesiones para ver todos los cierres anteriores.")

    # ══════════════════════════════════════════════════════════════════════════
    # MÓDULO 6 — COMPRAS
    # ══════════════════════════════════════════════════════════════════════════
    add_module_heading(doc, 6, "Gestión de Compras",
                       "Para clínicas que gestionan inventario o pagos a proveedores.")

    add_section_heading(doc, "6.1  Registrar un Proveedor")
    add_step(doc, 1, "Ir a Compras → Proveedores", [])
    add_step(doc, 2, "Nuevo proveedor", [
        "Haz clic en Nuevo proveedor.",
        "Ingresa nombre, contacto, email y RUT/documento fiscal.",
        "Haz clic en Guardar.",
    ])

    add_section_heading(doc, "6.2  Flujo de Compra (Presupuesto → Orden → Factura → Pago)")
    doc.add_paragraph(
        "El flujo de compras es simétrico al de ventas, pero del lado del proveedor."
    )
    add_step(doc, 1, "Crear presupuesto de compra", [
        "Ve a Compras → Presupuestos → Nuevo presupuesto.",
        "Selecciona el proveedor y agrega los ítems/productos.",
    ])
    add_step(doc, 2, "Convertir a orden de compra", [
        "Cuando el proveedor confirme, cambia el estado a Aceptado y crea la orden.",
    ])
    add_step(doc, 3, "Registrar la factura del proveedor", [
        "Cuando recibas la factura, ve a Compras → Facturas y regístrala.",
    ])
    add_step(doc, 4, "Registrar el pago al proveedor", [
        "Desde la factura de compra, haz clic en Registrar pago y completa los datos.",
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # MÓDULO 7 — ALERTAS Y COMUNICACIONES
    # ══════════════════════════════════════════════════════════════════════════
    add_module_heading(doc, 7, "Alertas y Comunicaciones Automáticas",
                       "Configura recordatorios automáticos para pacientes y staff.")

    add_section_heading(doc, "7.1  Configurar Plantillas de Comunicación")
    doc.add_paragraph(
        "Las plantillas definen el contenido de los mensajes automáticos (email o SMS) "
        "que el sistema envía a los pacientes."
    )
    add_step(doc, 1, "Ir a Sistema → Plantillas de comunicación", [])
    add_step(doc, 2, "Crear o editar una plantilla", [
        "Selecciona el tipo (Email o SMS) y el evento (ej. recordatorio de cita, factura vencida).",
        "Redacta el mensaje usando variables dinámicas como {paciente_nombre}, {fecha_cita}, {monto}.",
        "Haz clic en Guardar.",
    ])
    add_tip(doc, "Prueba la plantilla enviándote un mensaje de prueba antes de activarla.")

    add_section_heading(doc, "7.2  Configurar Reglas de Alertas")
    doc.add_paragraph(
        "Las reglas definen cuándo se dispara una alerta y qué acción tomar."
    )
    add_step(doc, 1, "Ir a Sistema → Reglas de alertas", [])
    add_step(doc, 2, "Nueva regla", [
        "Haz clic en Nueva regla.",
        "Define el evento disparador (ej. 'Factura pendiente hace 30 días').",
        "Selecciona la acción (ej. 'Enviar email con plantilla X al paciente').",
        "Activa la regla y haz clic en Guardar.",
    ])
    add_tip(doc, "Ejemplos útiles: recordatorio de cita 24 hs antes, aviso de factura vencida, bienvenida a nuevo paciente.")

    # ══════════════════════════════════════════════════════════════════════════
    # MÓDULO 8 — ESTUDIOS MÉDICOS
    # ══════════════════════════════════════════════════════════════════════════
    add_module_heading(doc, 8, "Estudios Médicos e Imágenes DICOM",
                       "Gestión y visualización de imágenes radiológicas.")

    add_section_heading(doc, "8.1  Subir y Visualizar un Estudio")
    add_step(doc, 1, "Ir a Estudios", ["En el menú lateral, haz clic en Estudios."])
    add_step(doc, 2, "Subir un estudio", [
        "Haz clic en Nuevo estudio.",
        "Selecciona el paciente.",
        "Sube el archivo DICOM (radiografía, tomografía, etc.).",
        "Agrega una descripción y la fecha del estudio.",
        "Haz clic en Guardar.",
    ])
    add_step(doc, 3, "Visualizar", [
        "Haz clic en el estudio para abrirlo en el visor integrado.",
        "Usa las herramientas del visor para ajustar brillo, contraste y hacer zoom.",
    ])
    add_step(doc, 4, "Compartir con el paciente", [
        "Ve a Estudios Compartidos.",
        "Selecciona el estudio y haz clic en Compartir.",
        "El paciente recibirá un enlace para ver el estudio desde su dispositivo.",
    ])
    add_tip(doc, "El visor de estudios es compatible con el estándar DICOM y permite visualizar series de imágenes.")

    # ══════════════════════════════════════════════════════════════════════════
    # MÓDULO 9 — TV DISPLAY
    # ══════════════════════════════════════════════════════════════════════════
    add_module_heading(doc, 9, "Pantalla de Sala de Espera (TV Display)",
                       "Proyecta los turnos del día en un monitor de sala de espera en tiempo real.")

    doc.add_paragraph(
        "La funcionalidad TV Display permite mostrar en una TV o monitor secundario los turnos "
        "del día agrupados por consultorio. Se controla desde el sistema sin necesidad de tocar "
        "la pantalla de TV."
    )

    add_section_heading(doc, "9.1  Configuración")
    add_step(doc, 1, "Ir a Pantalla TV", [
        "En el menú lateral, haz clic en Pantalla TV (ícono de monitor).",
    ])
    add_step(doc, 2, "Seleccionar consultorios", [
        "En el campo Consultorios, selecciona los calendarios que quieres mostrar en pantalla.",
        "Cada calendario aparecerá como una columna independiente en la TV.",
    ])
    add_step(doc, 3, "Configurar apariencia", [
        "Tema: elige entre Oscuro, Claro o Branded.",
        "Activa o desactiva el reloj y la fecha en el encabezado.",
        "Elige qué datos del paciente mostrar: nombre, doctor, hora del turno, próximo paciente.",
    ])
    add_step(doc, 4, "Información de la clínica", [
        "El nombre y logo se toman automáticamente de la configuración del sistema.",
        "Activa o desactiva la visibilidad del teléfono, dirección y email de la clínica en el encabezado.",
    ])
    add_step(doc, 5, "Videos", [
        "Videos de Promoción: URLs de videos o YouTube que se reproducen en pantalla completa entre cambios de paciente. Se reproducen en bucle.",
        "Columna de Videos Lateral: videos permanentes en una columna a la izquierda, derecha, arriba o abajo de la grilla. Elige 'Ninguna' para desactivar.",
    ])
    add_step(doc, 6, "Intervalo de actualización", [
        "Define cada cuántos minutos el sistema recarga los turnos del día automáticamente.",
    ])
    add_tip(doc, "El preview en la página de configuración refleja en tiempo real todos los cambios de configuración antes de encender la pantalla.")

    add_section_heading(doc, "9.2  Encender la Pantalla")
    add_step(doc, 1, "Hacer clic en 'Encender pantalla'", [
        "Se abrirá una nueva ventana del navegador con la pantalla TV en modo fullscreen.",
        "Arrastra esa ventana al monitor secundario o TV conectada a la computadora.",
    ])
    add_step(doc, 2, "Sincronización automática", [
        "La pantalla TV se sincroniza automáticamente con el sistema vía BroadcastChannel.",
        "No es necesario recargarla manualmente cuando se realizan cambios.",
    ])
    add_tip(doc, "Mantén la ventana TV abierta en el navegador del monitor de sala de espera durante toda la jornada.")

    add_section_heading(doc, "9.3  Widget Flotante de Control Rápido")
    doc.add_paragraph(
        "Una vez encendida, aparece un widget flotante en la barra del sistema (ícono de TV "
        "con indicador verde = encendida / rojo = apagada). Al hacer clic se despliega el menú de control:"
    )
    actions = [
        ("Siguiente paciente por consultorio", "Un botón por cada consultorio activo. Al presionarlo, la TV muestra una tarjeta de anuncio con el próximo paciente y avanza el turno."),
        ("Pausar pantalla", "Congela la visualización con superposición 'PAUSADO'. Útil durante recesos."),
        ("Mostrar promoción", "Activa el modo promo: los videos de promoción se reproducen en pantalla completa."),
        ("Actualizar datos", "Fuerza una recarga inmediata de los turnos desde el servidor."),
        ("Apagar", "Apaga la pantalla TV (muestra superposición 'APAGADO')."),
    ]
    for action, desc in actions:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(action + ": ")
        run.bold = True
        p.add_run(desc)
    doc.add_paragraph()

    add_section_heading(doc, "9.4  Flujo Típico de Uso Diario")
    add_step(doc, 1, "Inicio del día", [
        "Ir a Pantalla TV → verificar que los consultorios correctos están seleccionados.",
        "Hacer clic en Encender pantalla y mover la ventana al monitor de sala de espera.",
    ])
    add_step(doc, 2, "Durante la jornada", [
        "Cuando un paciente pasa a consultorio: abrir el widget flotante → hacer clic en 'Siguiente: [Consultorio X]'.",
        "La TV muestra el anuncio con el nombre del paciente entrante y avanza al siguiente turno.",
    ])
    add_step(doc, 3, "Fin del día o recesos", [
        "Pausar o apagar la pantalla desde el widget flotante.",
    ])
    add_tip(doc, "Si hay videos de promoción configurados, puedes activar el modo Promo durante esperas largas o entre pacientes.")

    # ══════════════════════════════════════════════════════════════════════════
    # APÉNDICE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading("Apéndice — Catálogo Clínico", level=1)
    doc.add_paragraph(
        "Antes de registrar historias clínicas, asegúrate de tener cargado el catálogo clínico. "
        "Estos datos maestros se configuran una sola vez en el menú Catálogo Clínico:"
    )
    catalog_items = [
        ("Afecciones / Patologías", "Listado de diagnósticos y condiciones médicas. Ej: 'Caries', 'Gingivitis', 'Bruxismo'."),
        ("Medicamentos", "Fármacos que pueden prescribirse. Incluye nombre comercial y genérico."),
        ("Condiciones dentales", "Estados de las piezas dentales para el odontograma. Ej: 'Corona', 'Extracción', 'Sellante'."),
        ("Superficies dentales", "Nomenclatura de superficies para documentar tratamientos en piezas específicas."),
    ]
    for name, desc in catalog_items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(name + ": ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph()
    doc.add_heading("Glosario Rápido", level=1)
    terms = [
        ("Presupuesto", "Propuesta de tratamiento con precios enviada al paciente."),
        ("Orden", "Confirmación interna de un presupuesto aceptado; autoriza el trabajo."),
        ("Factura", "Documento de cobro formal generado a partir de una orden."),
        ("Sesión de caja", "Turno de trabajo en caja con apertura y cierre de balance."),
        ("DICOM", "Formato estándar para imágenes médicas digitales (radiografías, etc.)."),
        ("Rol", "Conjunto de permisos asignado a un usuario del sistema."),
    ]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["Término", "Definición"]):
        hdr[i].paragraphs[0].add_run(h).bold = True
        set_cell_bg(hdr[i], "DBEAFE")
    for term, definition in terms:
        row = tbl.add_row().cells
        row[0].paragraphs[0].add_run(term).bold = True
        row[1].paragraphs[0].add_run(definition)

    # ── Pie de página ──────────────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Invoke IA  ·  Guía de Usuario  ·  2026")
        run.font.size = Pt(8)
        run.font.color.rgb = COLOR_GRAY

    # ── Guardar ────────────────────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"✓ Documento generado: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
